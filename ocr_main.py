# coding=utf-8
import os
import time
import hashlib
import base64
import requests
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

import json

# ========== 配置区 =================================================
URL = "http://webapi.xfyun.cn/v1/service/v1/ocr/handwriting"
APPID = ""
API_KEY = ""
language = "cn|en"
location = "false"
# ==================================================================


def getHeader():
    curTime = str(int(time.time()))
    param = json.dumps({"language": language, "location": location})
    paramBase64 = base64.b64encode(param.encode('utf-8')).decode('utf-8')
    checkSum_str = API_KEY + curTime + paramBase64
    checkSum = hashlib.md5(checkSum_str.encode('utf-8')).hexdigest()
    header = {
        'X-CurTime': curTime,
        'X-Param': paramBase64,
        'X-Appid': APPID,
        'X-CheckSum': checkSum,
        'Content-Type': 'application/x-www-form-urlencoded; charset=utf-8',
    }
    return header


def ocr_and_extract_text(image_path):
    with open(image_path, 'rb') as f:
        imgfile = f.read()
    data = {'image': base64.b64encode(imgfile).decode('utf-8')}
    headers = getHeader()
    resp = requests.post(URL, headers=headers, data=data)
    result = resp.json()

    if result.get('code') != '0':
        raise RuntimeError(f"OCR 失败: {result.get('desc')}")

    try:
        blocks = result.get('data', {}).get('block', [])
        if not blocks:
            raise ValueError("无识别结果（block 为空）")

        text_list = []
        for block in blocks:
            lines = block.get('line', [])
            for line in lines:
                words = line.get('word', [])
                for word in words:
                    content = word.get('content', '')
                    text_list.append(content)
        return ''.join(text_list)
    except Exception as e:
        raise RuntimeError(f"OCR 返回结构错误：{e}")


def create_word(doc_path, all_text_blocks,folder_display_name):

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)

    for text in all_text_blocks:
        para_before = doc.add_paragraph("修改前：")
        para_before.paragraph_format.space_before = Pt(0)
        para_before.paragraph_format.space_after = Pt(0)
        para_before.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
        para_before.paragraph_format.line_spacing = Pt(12)

        # 添加 ——姓名（文件夹名）
        para_name = doc.add_paragraph(f"——{folder_display_name}")
        para_name.paragraph_format.space_before = Pt(0)
        para_name.paragraph_format.space_after = Pt(0)
        para_name.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
        para_name.paragraph_format.line_spacing = Pt(12)
        para_name.alignment = WD_ALIGN_PARAGRAPH.CENTER



        para_text = doc.add_paragraph(text)
        para_text.paragraph_format.space_before = Pt(0)
        para_text.paragraph_format.space_after = Pt(0)
        para_text.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
        para_text.paragraph_format.line_spacing = Pt(12)
        para_text.paragraph_format.first_line_indent = Cm(0.74)

        doc.add_page_break()

        para_after = doc.add_paragraph("修改后：")
        para_after.paragraph_format.space_before = Pt(0)
        para_after.paragraph_format.space_after = Pt(0)
        para_after.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
        para_after.paragraph_format.line_spacing = Pt(12)
        para_text.paragraph_format.first_line_indent = Cm(0.74)

    doc.save(doc_path)


def has_images(folder_path):
    return any(
        f.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp'))
        for f in os.listdir(folder_path)
    )


def process_folder(folder_path):
    folder_name = os.path.basename(folder_path)
    all_text_blocks = []

    for filename in os.listdir(folder_path):
        if filename.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp')):
            img_path = os.path.join(folder_path, filename)
            try:
                print(f"正在识别：{img_path}")
                text = ocr_and_extract_text(img_path)
                all_text_blocks.append(text)  # 应使用 append
            except Exception as e:
                print(f"识别失败：{img_path}，原因：{e}")

    if all_text_blocks:
        full_text = '\n'.join(all_text_blocks)  # 合并为一个整体文本
        doc_path = os.path.join(folder_path, f"{folder_name}.docx")
        print(f"正在生成 Word：{doc_path}")
        create_word(doc_path, [full_text],folder_name)
    else:
        print(f"{folder_path} 中没有可识别的图片")


def process_all(root_dir):
    if has_images(root_dir):
        process_folder(root_dir)
    else:
        for sub in os.listdir(root_dir):
            sub_path = os.path.join(root_dir, sub)
            if os.path.isdir(sub_path) and has_images(sub_path):
                process_folder(sub_path)


if __name__ == '__main__':
    ROOT_DIR = input("请输入要处理的文件夹路径：").strip('" ')
    if not os.path.isdir(ROOT_DIR):
        print("无效路径！请确认后再试。")
    else:
        process_all(ROOT_DIR)
        print("全部处理完成！")
