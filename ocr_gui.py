import base64
import json
import shutil
import threading
import traceback
from concurrent.futures import FIRST_COMPLETED, ThreadPoolExecutor, wait
from datetime import datetime
import sys
import os
import re
from pathlib import Path
from copy import deepcopy

from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QLineEdit, QPushButton, QCheckBox, QComboBox, QTextEdit,
    QScrollArea, QFrame, QTableWidget, QTableWidgetItem, QHeaderView,
    QAbstractItemView, QSplitter, QGroupBox, QFormLayout, QSizePolicy,
    QFileDialog, QMessageBox,
)
from PySide6.QtCore import Qt, Signal, QObject, QTimer
from PySide6.QtGui import QFont, QColor, QIcon

from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_LINE_SPACING, WD_BREAK, WD_ALIGN_PARAGRAPH
from openai import OpenAI
from llm_client import fetch_models


# ================= 默认配置 =================
# 打包后 _MEIPASS 是临时目录，EXE 所在目录才是用户目录
def _exe_dir() -> Path:
    if getattr(sys, "frozen", False):
        return Path(sys.executable).resolve().parent
    return Path(__file__).resolve().parent

PERSONAL_CONFIG = Path(r"D:\person_data\ocer助手\presson.json")
LOCAL_CONFIG = _exe_dir() / "config.json"
PROJECT_DIR = Path(__file__).resolve().parent
CRASH_LOG = _exe_dir() / "ocr_gui_crash.log"
STARTUP_LOG = _exe_dir() / "ocr_gui_startup.log"

DEFAULT_CONFIG = {
    "OCR": {
        "PROVIDER": "xfyun_handwriting",
        "XFYUN": {
            "URL": "http://webapi.xfyun.cn/v1/service/v1/ocr/handwriting",
            "APPID": "",
            "API_KEY": "",
            "LANGUAGE": "cn|en",
            "LOCATION": "false",
        },
    },
    "LLM": {
        "PROVIDERS": {
            "deepseek": {"API_KEY": "", "MODEL": "deepseek-chat", "BASE_URL": "https://api.deepseek.com/v1"},
            "openai": {"API_KEY": "", "MODEL": "gpt-4o-mini", "BASE_URL": "https://api.openai.com/v1"},
            "custom": {"API_KEY": "", "MODEL": "", "BASE_URL": ""},
        },
        "TASKS": {
            "typo_fix": {"ENABLED": False, "PROVIDER": "deepseek", "PROMPT": "{text}"},
            "editor": {"ENABLED": False, "PROVIDER": "deepseek", "PROMPT": "{text}", "COUNT_MIN": None, "COUNT_MAX": None},
        },
    },
    "APP": {"ROOT_DIR": "", "DEBUG": False},
}


def _is_bad_marshal_error(exc: BaseException) -> bool:
    return "bad marshal data" in str(exc).lower()


def _clear_project_bytecode_caches():
    """Remove local bytecode caches that can break after switching Python versions."""
    if getattr(sys, "frozen", False):
        return
    for cache_dir in PROJECT_DIR.rglob("__pycache__"):
        try:
            shutil.rmtree(cache_dir)
        except OSError:
            pass


def _format_exception_for_log(exc: BaseException) -> str:
    detail = "".join(traceback.format_exception(type(exc), exc, exc.__traceback__)).strip()
    if _is_bad_marshal_error(exc):
        detail += "\n\n提示：bad marshal data 通常是 Python 字节码缓存（.pyc）版本不匹配或损坏导致的。程序已尝试清理项目内 __pycache__，如仍失败请重新打包/重启程序。"
    return detail


def _write_crash_log(exc_type, exc, tb):
    try:
        with CRASH_LOG.open("a", encoding="utf-8") as f:
            f.write(f"\n[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] 未捕获异常\n")
            f.write("".join(traceback.format_exception(exc_type, exc, tb)))
            f.write("\n")
    except Exception:
        pass


def _write_startup_log(message: str):
    try:
        with STARTUP_LOG.open("a", encoding="utf-8") as f:
            f.write(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] {message}\n")
    except Exception:
        pass


sys.excepthook = _write_crash_log

from config_migrate import ensure_new_schema


def load_config(path: Path = None):
    """加载配置：优先个人目录 → 当前目录 → 默认"""
    if path:
        cfg_path = Path(path)
    elif PERSONAL_CONFIG.exists():
        cfg_path = PERSONAL_CONFIG
    elif LOCAL_CONFIG.exists():
        cfg_path = LOCAL_CONFIG
    else:
        return deepcopy(DEFAULT_CONFIG)
    try:
        with cfg_path.open("r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return deepcopy(DEFAULT_CONFIG)


def save_config(config, path: Path = None):
    """保存配置：优先保存到个人目录，不存在则保存到当前目录"""
    if path:
        cfg_path = Path(path)
    elif PERSONAL_CONFIG.parent.exists():
        cfg_path = PERSONAL_CONFIG
    else:
        cfg_path = LOCAL_CONFIG
    cfg_path.parent.mkdir(parents=True, exist_ok=True)
    with cfg_path.open("w", encoding="utf-8") as f:
        json.dump(config, f, indent=2, ensure_ascii=False)


# ================= Provider helpers =================
API_PROVIDER_PRESETS = {"deepseek": "https://api.deepseek.com/v1", "openai": "https://api.openai.com/v1"}


def _normalize_provider_name(name: str) -> str:
    return (name or "").strip().lower()


def _ensure_provider_exists(config, name: str):
    p_name = _normalize_provider_name(name)
    if not p_name:
        return ""
    config.setdefault("LLM", {}).setdefault("PROVIDERS", {}).setdefault(p_name, {})
    p_cfg = config["LLM"]["PROVIDERS"][p_name]
    p_cfg.setdefault("API_KEY", "")
    p_cfg.setdefault("BASE_URL", API_PROVIDER_PRESETS.get(p_name, ""))
    p_cfg.setdefault("MODEL", "deepseek-chat" if p_name == "deepseek" else "gpt-4o-mini")
    return p_name


def _provider_name_list(config):
    providers = (config.get("LLM", {}) or {}).get("PROVIDERS", {}) or {}
    names = {k for k in providers.keys() if isinstance(k, str) and k.strip()}
    names.update(API_PROVIDER_PRESETS.keys())
    return sorted(names)


# ================= File helpers =================
def iter_files_limited(folder, max_depth=4):
    folder = os.path.abspath(folder)
    for root_dir, dirs, files in os.walk(folder, topdown=True):
        rel = os.path.relpath(root_dir, folder)
        depth = 0 if rel == os.curdir else len(rel.split(os.sep))
        if depth >= max_depth - 1:
            dirs[:] = []
        yield root_dir, files


def has_images_folder(path: str) -> bool:
    try:
        return any(
            f.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp'))
            for f in os.listdir(path)
            if os.path.isfile(os.path.join(path, f))
        )
    except Exception:
        return False


def scan_folder_for_tasks(folder: str) -> list[str]:
    folder = os.path.abspath(folder)
    tasks = []
    if has_images_folder(folder):
        tasks.append(folder)
    try:
        for name in sorted(os.listdir(folder)):
            if name == "旧":
                continue
            child = os.path.join(folder, name)
            if os.path.isdir(child) and has_images_folder(child):
                tasks.append(child)
    except Exception:
        pass
    return tasks


def infer_student_and_essay(folder_name: str):
    if "_" in folder_name:
        parts = folder_name.split("_", 1)
        return parts[0], parts[1]
    if "-" in folder_name:
        parts = folder_name.split("-", 1)
        return parts[0], parts[1]
    return folder_name, folder_name


def count_chinese_characters(text: str) -> int:
    return sum(1 for ch in text if not ch.isspace())


def determine_word_count_bounds(original_count: int):
    if original_count >= 850:
        return max(700, original_count - 30), original_count + 30
    if original_count >= 800:
        return 820, 850
    return 700, 820


# ================= Log signal =================
class LogSignal(QObject):
    log_message = Signal(str)
    task_status = Signal(str, str, str, str, str, str, str, str, str)  # +essay_title
    ai_log_message = Signal(str)
    ai_tasks_loaded = Signal(list)
    ai_task_status = Signal(str, str, str, str, str)
    doc_ai_log_message = Signal(str)
    doc_ai_tasks_loaded = Signal(list)
    doc_ai_task_status = Signal(str, str, str)  # task_path, status, log_msg


# ================= Collapsible Section =================
class CollapsibleSection(QWidget):
    def __init__(self, title="", collapsed=True, parent=None):
        super().__init__(parent)
        self._collapsed = collapsed
        self._title = title

        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        self.toggle_btn = QPushButton(("▶ " + title) if collapsed else ("▼ " + title))
        self.toggle_btn.setStyleSheet(
            "QPushButton { text-align: left; border: none; background: transparent;"
            "color: #1a73e8; font-weight: bold; font-size: 13px; padding: 6px 8px; }"
            "QPushButton:hover { background: #e8f0fe; }"
        )
        self.toggle_btn.clicked.connect(self._toggle)
        layout.addWidget(self.toggle_btn)

        self.content_widget = QWidget()
        self.content_layout = QVBoxLayout(self.content_widget)
        self.content_layout.setContentsMargins(8, 4, 8, 8)
        layout.addWidget(self.content_widget)

        if collapsed:
            self.content_widget.setVisible(False)

    def _toggle(self):
        self._collapsed = not self._collapsed
        self.content_widget.setVisible(not self._collapsed)
        prefix = "▼ " if not self._collapsed else "▶ "
        self.toggle_btn.setText(prefix + self._title)

    def add_widget(self, w):
        self.content_layout.addWidget(w)

    def add_layout(self, l):
        self.content_layout.addLayout(l)


# ================= Main Window =================
class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        _clear_project_bytecode_caches()
        self.config = ensure_new_schema(load_config())
        self.hidden_api_keys = {}
        self.task_queue = []
        self.completed_tasks = set()  # 已完成的任务路径集合
        self.finished_tasks = set()  # 当前运行批次中已经尝试过的任务路径集合
        self.in_progress_tasks = set()  # 当前正在处理的任务路径集合
        self.is_processing = False
        self.max_parallel_tasks = 3
        self.queue_lock = threading.Lock()
        self.log_signal = LogSignal()
        self.log_signal.log_message.connect(self._append_log)
        self.log_signal.task_status.connect(self._update_task_status)
        self.log_signal.ai_tasks_loaded.connect(self._render_ai_queue)
        self.log_signal.ai_log_message.connect(self._append_log_ai)
        self.log_signal.ai_task_status.connect(self._update_ai_task_status)
        self.log_signal.doc_ai_log_message.connect(self._append_log_doc_ai)
        self.log_signal.doc_ai_tasks_loaded.connect(self._render_doc_ai_queue)
        self.log_signal.doc_ai_task_status.connect(self._update_doc_ai_task_status)

        self.setWindowTitle("Composition OCR Assistant 作文修改助手 v1.2")
        self.resize(1100, 800)

        # 启用拖放
        self.setAcceptDrops(True)

        # Set icon
        for ico_name in ("app.ico",):
            ico_path = Path(__file__).resolve().parent / ico_name
            if ico_path.exists():
                self.setWindowIcon(QIcon(str(ico_path)))
                break

        # Central widget
        central = QWidget()
        self.setCentralWidget(central)
        main_layout = QVBoxLayout(central)
        main_layout.setContentsMargins(8, 8, 8, 8)
        main_layout.setSpacing(4)

        # Top bar
        top_bar = QHBoxLayout()
        top_bar.addWidget(QLabel("功能选择:"))
        self.btn_ocr = QPushButton("图片转作文")
        self.btn_ocr.setFixedWidth(120)
        self.btn_ai = QPushButton("文档作文处理")
        self.btn_ai.setFixedWidth(120)
        self.btn_doc_ai = QPushButton("文档ai识别处理")
        self.btn_doc_ai.setFixedWidth(140)
        self.btn_config = QPushButton("配置编辑")
        self.btn_config.setFixedWidth(100)
        self.btn_config.clicked.connect(self._open_config_editor)
        top_bar.addWidget(self.btn_ocr)
        top_bar.addWidget(self.btn_ai)
        top_bar.addWidget(self.btn_doc_ai)
        top_bar.addStretch()
        top_bar.addWidget(self.btn_config)
        main_layout.addLayout(top_bar)

        # Page stack
        self.page_ocr = QWidget()
        self.page_ai = QWidget()
        self.page_doc_ai = QWidget()
        self.page_ocr.hide()
        self.page_ai.hide()
        self.page_doc_ai.hide()
        main_layout.addWidget(self.page_ocr)
        main_layout.addWidget(self.page_ai)
        main_layout.addWidget(self.page_doc_ai)

        self.btn_ocr.clicked.connect(lambda: self._show_page("ocr"))
        self.btn_ai.clicked.connect(lambda: self._show_page("ai"))
        self.btn_doc_ai.clicked.connect(lambda: self._show_page("doc_ai"))

        self._init_page_ocr()
        self._init_page_ai()
        self._init_page_doc_ai()
        self._show_page("ocr")

    def _show_page(self, name):
        self.page_ocr.hide()
        self.page_ai.hide()
        self.page_doc_ai.hide()
        if name == "ocr":
            self.page_ocr.show()
        elif name == "ai":
            self.page_ai.show()
        else:
            self.page_doc_ai.show()

    # 拖放支持
    def dragEnterEvent(self, event):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()

    def _make_key_toggle(self, entry: QLineEdit, layout: QHBoxLayout):
        """在 layout 中 entry 后面添加一个眼睛按钮，点击切换密码显示/隐藏"""
        btn = QPushButton("👁")
        btn.setFixedWidth(28)
        btn.setCheckable(True)
        btn.setStyleSheet("QPushButton { border: none; font-size: 14px; }")
        btn.toggled.connect(lambda checked: entry.setEchoMode(QLineEdit.Normal if checked else QLineEdit.Password))
        layout.addWidget(btn)
        return btn

    def dropEvent(self, event):
        urls = event.mimeData().urls()
        if not urls:
            return
        
        dropped_paths = [url.toLocalFile() for url in urls]
        
        # 分类拖入的内容
        docx_files = []
        folders = []
        image_files = []
        
        for path in dropped_paths:
            if os.path.isdir(path):
                folders.append(path)
            elif path.lower().endswith('.docx') and not os.path.basename(path).startswith('~$'):
                # 检查文件是否在"修改后"文件夹下
                if "修改后" not in path.split(os.sep):
                    docx_files.append(path)
            elif path.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp')):
                image_files.append(path)
        
        # 根据当前页面和拖入内容决定处理方式
        if docx_files or folders:
            # 切换到"文档ai识别处理"页面
            self._show_page("doc_ai")
            
            # 如果拖入的是文件夹，设置路径
            if folders:
                self.doc_ai_path_entry.setText(folders[0])
                self.log_signal.doc_ai_log_message.emit(f"已设置文件夹路径：{folders[0]}")
            
            # 如果拖入的是docx文件，直接开始处理
            if docx_files:
                self._process_dropped_docx_files(docx_files)
        
        elif image_files:
            # 切换到"图片转作文"页面
            self._show_page("ocr")
            self.log_signal.log_message.emit(f"拖入了 {len(image_files)} 个图片文件")

    def _process_dropped_docx_files(self, docx_files):
        """处理拖入的docx文件"""
        # 检查AI配置
        api_key = self.doc_ai_key_entry.text().strip()
        if not api_key:
            self.log_signal.doc_ai_log_message.emit("请先填写 API Key")
            return
        
        # 将拖入的文件添加到任务列表
        added_count = self._add_doc_ai_tasks(docx_files)
        self.log_signal.doc_ai_log_message.emit(f"添加了 {added_count} 个docx文件到任务列表")
        
        # 在新线程中处理
        threading.Thread(target=self._run_dropped_docx_workflow, args=(docx_files,), daemon=True).start()

    def _run_dropped_docx_workflow(self, docx_files):
        """处理拖入的docx文件的工作流"""
        import shutil
        
        api_key = self.doc_ai_key_entry.text().strip()
        base_url = self.doc_ai_url_entry.text().strip()
        model = self.doc_ai_model_combo.currentText().strip() or "deepseek-chat"
        max_parallel = int(self.doc_ai_parallel_spin.currentText())

        # 保存配置
        selected_provider = _ensure_provider_exists(self.config, self.doc_ai_provider_combo.currentText() or "deepseek")
        cfg = self.config
        cfg.setdefault("LLM", {})
        cfg["LLM"].setdefault("PROVIDERS", {})
        cfg["LLM"]["PROVIDERS"].setdefault(selected_provider, {})
        cfg["LLM"]["PROVIDERS"][selected_provider]["API_KEY"] = api_key
        cfg["LLM"]["PROVIDERS"][selected_provider]["BASE_URL"] = base_url
        cfg["LLM"]["PROVIDERS"][selected_provider]["MODEL"] = model
        save_config(cfg)

        # 获取自定义提示词
        custom_prompt = self.doc_ai_prompt_text.toPlainText().strip()
        if not custom_prompt:
            custom_prompt = "下面是一篇中文文章，请你【只修改错别字和明显的识别错误】。\n要求：1. 不改变原意 2. 不润色文风 3. 不增删内容 4. 保持原有段落结构 5. 只输出修改后的完整文章正文\n"

        # 获取字数限制
        count_min = None
        count_max = None
        min_text = self.doc_ai_count_min.text().strip()
        max_text = self.doc_ai_count_max.text().strip()
        if min_text:
            try:
                count_min = int(min_text)
            except ValueError:
                pass
        if max_text:
            try:
                count_max = int(max_text)
            except ValueError:
                pass

        # 创建AI客户端
        try:
            client = OpenAI(api_key=api_key, base_url=base_url)
        except Exception as e:
            self.log_signal.doc_ai_log_message.emit(f"创建AI客户端失败: {e}")
            return

        # 使用线程池处理文件
        with ThreadPoolExecutor(max_workers=max_parallel) as executor:
            futures = []
            for docx_path in docx_files:
                future = executor.submit(self._process_single_docx, docx_path, client, model, custom_prompt, count_min, count_max)
                futures.append((future, docx_path))

            # 等待所有任务完成
            for future, docx_path in futures:
                try:
                    future.result()
                except Exception as e:
                    self.log_signal.doc_ai_log_message.emit(f"  {os.path.basename(docx_path)}: 线程异常 - {e}")
                    self.log_signal.doc_ai_task_status.emit(docx_path, "失败", str(e))

        self.log_signal.doc_ai_log_message.emit("处理完成！")

    def _append_log(self, msg):
        ts = datetime.now().strftime("%H:%M:%S")
        self.log_text.append(f"[{ts}] {msg}")

    def _append_log_ai(self, msg):
        ts = datetime.now().strftime("%H:%M:%S")
        self.ai_log_text.append(f"[{ts}] {msg}")

    def _append_log_doc_ai(self, msg):
        ts = datetime.now().strftime("%H:%M:%S")
        self.doc_ai_log_text.append(f"[{ts}] {msg}")

    def _open_config_editor(self):
        from config_editor_ui import open_config_editor_form
        # 确定实际配置文件路径
        if PERSONAL_CONFIG.exists():
            cfg_file = PERSONAL_CONFIG
        elif LOCAL_CONFIG.exists():
            cfg_file = LOCAL_CONFIG
        else:
            cfg_file = PERSONAL_CONFIG  # 默认保存到个人目录
        open_config_editor_form(
            parent=self,
            config=self.config,
            config_file=cfg_file,
            hidden_api_keys=self.hidden_api_keys,
            on_saved=self._on_config_saved,
        )

    def _on_config_saved(self, new_cfg):
        self.config = ensure_new_schema(new_cfg)
        # Refresh OCR fields
        self.url_entry.setText(self.config.get("OCR", {}).get("XFYUN", {}).get("URL", ""))
        self.appid_entry.setText(self.config.get("OCR", {}).get("XFYUN", {}).get("APPID", ""))
        self.path_entry.setText(self.config.get("APP", {}).get("ROOT_DIR", ""))
        # Refresh checkboxes
        self.use_typo_fix.setChecked(bool((self.config.get("LLM", {}).get("TASKS", {}).get("typo_fix", {}) or {}).get("ENABLED", False)))
        self.typo_prompt_text.setPlainText((self.config.get("LLM", {}).get("TASKS", {}).get("typo_fix", {}) or {}).get("PROMPT", "{text}"))
        self.use_editor.setChecked(bool((self.config.get("LLM", {}).get("TASKS", {}).get("editor", {}) or {}).get("ENABLED", False)))
        self.editor_prompt_text.setPlainText((self.config.get("LLM", {}).get("TASKS", {}).get("editor", {}) or {}).get("PROMPT", "{text}"))
        # Refresh provider combos
        names = _provider_name_list(self.config)
        for combo in [self.typo_provider_combo, self.editor_provider_combo]:
            current = combo.currentText()
            combo.clear()
            combo.addItems(names)
            idx = combo.findText(current)
            if idx >= 0:
                combo.setCurrentIndex(idx)

    # ===================== PAGE 1: OCR =====================
    def _init_page_ocr(self):
        layout = QVBoxLayout(self.page_ocr)
        layout.setContentsMargins(0, 0, 0, 0)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll_widget = QWidget()
        scroll_layout = QVBoxLayout(scroll_widget)
        scroll_layout.setContentsMargins(4, 4, 4, 4)
        scroll_layout.setSpacing(6)

        # ---- Baidu correction (collapsible) ----
        baidu_sec = CollapsibleSection("百度图片矫正（OCR前自动矫正倾斜/弯曲文档）", collapsed=True)
        self.use_baidu_correction = QCheckBox("启用图片矫正（去阴影+透视变换）")
        self.use_baidu_correction.setChecked(bool(self.config.get("OCR", {}).get("BAIDU_CORRECTION", {}).get("ENABLED", False)))
        baidu_sec.add_widget(self.use_baidu_correction)

        baidu_key_layout = QHBoxLayout()
        baidu_key_layout.addWidget(QLabel("百度 API Key"))
        self.baidu_api_key_entry = QLineEdit(self.config.get("OCR", {}).get("BAIDU_CORRECTION", {}).get("API_KEY", ""))
        self.baidu_api_key_entry.setEchoMode(QLineEdit.Password)
        baidu_key_layout.addWidget(self.baidu_api_key_entry, 1)
        self._make_key_toggle(self.baidu_api_key_entry, baidu_key_layout)
        baidu_sec.add_layout(baidu_key_layout)

        baidu_secret_layout = QHBoxLayout()
        baidu_secret_layout.addWidget(QLabel("百度 Secret Key"))
        baidu_secret_key_entry = QLineEdit(self.config.get("OCR", {}).get("BAIDU_CORRECTION", {}).get("SECRET_KEY", ""))
        baidu_secret_key_entry.setEchoMode(QLineEdit.Password)
        baidu_secret_layout.addWidget(baidu_secret_key_entry, 1)
        self._make_key_toggle(baidu_secret_key_entry, baidu_secret_layout)
        baidu_sec.add_layout(baidu_secret_layout)
        self.baidu_secret_key_entry = baidu_secret_key_entry

        scroll_layout.addWidget(baidu_sec)

        # ---- OCR config (collapsible) ----
        ocr_sec = CollapsibleSection("OCR 识别配置", collapsed=True)

        ocr_url_layout = QHBoxLayout()
        ocr_url_layout.addWidget(QLabel("OCR 接口 URL"))
        self.url_entry = QLineEdit(self.config.get("OCR", {}).get("XFYUN", {}).get("URL", ""))
        ocr_url_layout.addWidget(self.url_entry, 1)
        ocr_sec.add_layout(ocr_url_layout)

        ocr_appid_layout = QHBoxLayout()
        ocr_appid_layout.addWidget(QLabel("APPID"))
        self.appid_entry = QLineEdit(self.config.get("OCR", {}).get("XFYUN", {}).get("APPID", ""))
        ocr_appid_layout.addWidget(self.appid_entry, 1)
        ocr_sec.add_layout(ocr_appid_layout)

        ocr_apikey_layout = QHBoxLayout()
        ocr_apikey_layout.addWidget(QLabel("API_KEY"))
        self.apikey_entry = QLineEdit(self.config.get("OCR", {}).get("XFYUN", {}).get("API_KEY", ""))
        self.apikey_entry.setEchoMode(QLineEdit.Password)
        ocr_apikey_layout.addWidget(self.apikey_entry, 1)
        self._make_key_toggle(self.apikey_entry, ocr_apikey_layout)
        ocr_sec.add_layout(ocr_apikey_layout)

        scroll_layout.addWidget(ocr_sec)

        # ---- AI typo fix ----
        typo_group = QGroupBox("第一步：AI 错别字修正")
        typo_layout = QVBoxLayout()

        # API Key
        row1 = QHBoxLayout()
        row1.addWidget(QLabel("AI API Key"))
        self.typo_api_key = QLineEdit()
        self.typo_api_key.setEchoMode(QLineEdit.Password)
        typo_provider = _normalize_provider_name(
            (self.config.get("LLM", {}).get("TASKS", {}).get("typo_fix", {}) or {}).get("PROVIDER", "deepseek")
        ) or "deepseek"
        self.typo_api_key.setText((self.config.get("LLM", {}).get("PROVIDERS", {}).get(typo_provider, {}) or {}).get("API_KEY", ""))
        row1.addWidget(self.typo_api_key, 1)
        self._make_key_toggle(self.typo_api_key, row1)
        self.use_typo_fix = QCheckBox("启用 AI 错别字自动修正（较慢）")
        self.use_typo_fix.setChecked(bool((self.config.get("LLM", {}).get("TASKS", {}).get("typo_fix", {}) or {}).get("ENABLED", False)))
        row1.addWidget(self.use_typo_fix)
        typo_layout.addLayout(row1)

        # Provider + Base URL
        row2 = QHBoxLayout()
        row2.addWidget(QLabel("AI Provider"))
        self.typo_provider_combo = QComboBox()
        provider_names = _provider_name_list(self.config)
        self.typo_provider_combo.addItems(provider_names)
        idx = self.typo_provider_combo.findText(typo_provider)
        if idx >= 0:
            self.typo_provider_combo.setCurrentIndex(idx)
        self.typo_provider_combo.currentTextChanged.connect(self._on_typo_provider_change)
        row2.addWidget(self.typo_provider_combo)
        btn_new_typo = QPushButton("+ 新增")
        btn_new_typo.setFixedWidth(68)
        btn_new_typo.clicked.connect(lambda: self._add_provider(self.typo_provider_combo))
        row2.addWidget(btn_new_typo)
        row2.addWidget(QLabel("Base URL"))
        self.typo_base_entry = QLineEdit((self.config.get("LLM", {}).get("PROVIDERS", {}).get(typo_provider, {}) or {}).get("BASE_URL", ""))
        row2.addWidget(self.typo_base_entry, 1)
        typo_layout.addLayout(row2)

        # Prompt
        row3 = QHBoxLayout()
        row3.addWidget(QLabel("自定义提示词"))
        self.typo_prompt_text = QTextEdit()
        self.typo_prompt_text.setMaximumHeight(120)
        self.typo_prompt_text.setPlainText(
            (self.config.get("LLM", {}).get("TASKS", {}).get("typo_fix", {}) or {}).get("PROMPT")
            or DEFAULT_CONFIG["LLM"]["TASKS"]["typo_fix"]["PROMPT"]
        )
        row3.addWidget(self.typo_prompt_text, 1)
        typo_layout.addLayout(row3)

        typo_group.setLayout(typo_layout)
        scroll_layout.addWidget(typo_group)

        # ---- Editor (step 2) ----
        editor_group = QGroupBox("第二步：AI 修改作文")
        editor_layout = QVBoxLayout()

        row4 = QHBoxLayout()
        self.use_editor = QCheckBox("启用 第二步 修改作文")
        self.use_editor.setChecked(bool((self.config.get("LLM", {}).get("TASKS", {}).get("editor", {}) or {}).get("ENABLED", False)))
        row4.addWidget(self.use_editor)
        editor_layout.addLayout(row4)

        row5 = QHBoxLayout()
        row5.addWidget(QLabel("AI API Key"))
        self.editor_api_key = QLineEdit()
        self.editor_api_key.setEchoMode(QLineEdit.Password)
        editor_provider = _normalize_provider_name(
            (self.config.get("LLM", {}).get("TASKS", {}).get("editor", {}) or {}).get("PROVIDER", "deepseek")
        ) or "deepseek"
        self.editor_api_key.setText((self.config.get("LLM", {}).get("PROVIDERS", {}).get(editor_provider, {}) or {}).get("API_KEY", ""))
        row5.addWidget(self.editor_api_key, 1)
        self._make_key_toggle(self.editor_api_key, row5)
        editor_layout.addLayout(row5)

        row6 = QHBoxLayout()
        row6.addWidget(QLabel("AI Provider"))
        self.editor_provider_combo = QComboBox()
        self.editor_provider_combo.addItems(provider_names)
        idx = self.editor_provider_combo.findText(editor_provider)
        if idx >= 0:
            self.editor_provider_combo.setCurrentIndex(idx)
        self.editor_provider_combo.currentTextChanged.connect(self._on_editor_provider_change)
        row6.addWidget(self.editor_provider_combo)
        btn_new_editor = QPushButton("+ 新增")
        btn_new_editor.setFixedWidth(68)
        btn_new_editor.clicked.connect(lambda: self._add_provider(self.editor_provider_combo))
        row6.addWidget(btn_new_editor)
        row6.addWidget(QLabel("Base URL"))
        self.editor_base_entry = QLineEdit((self.config.get("LLM", {}).get("PROVIDERS", {}).get(editor_provider, {}) or {}).get("BASE_URL", ""))
        row6.addWidget(self.editor_base_entry, 1)
        editor_layout.addLayout(row6)

        row7 = QHBoxLayout()
        row7.addWidget(QLabel("自定义提示词"))
        self.editor_prompt_text = QTextEdit()
        self.editor_prompt_text.setMaximumHeight(120)
        self.editor_prompt_text.setPlainText(
            (self.config.get("LLM", {}).get("TASKS", {}).get("editor", {}) or {}).get("PROMPT")
            or DEFAULT_CONFIG["LLM"]["TASKS"]["editor"]["PROMPT"]
        )
        row7.addWidget(self.editor_prompt_text, 1)
        editor_layout.addLayout(row7)

        row8 = QHBoxLayout()
        row8.addWidget(QLabel("目标字数"))
        self.editor_count_min = QLineEdit(str((self.config.get("LLM", {}).get("TASKS", {}).get("editor", {}) or {}).get("COUNT_MIN") or ""))
        self.editor_count_min.setFixedWidth(100)
        row8.addWidget(self.editor_count_min)
        row8.addWidget(QLabel("-"))
        self.editor_count_max = QLineEdit(str((self.config.get("LLM", {}).get("TASKS", {}).get("editor", {}) or {}).get("COUNT_MAX") or ""))
        self.editor_count_max.setFixedWidth(100)
        row8.addWidget(self.editor_count_max)
        row8.addWidget(QLabel("（空白表示自动）"))
        row8.addStretch()
        editor_layout.addLayout(row8)

        editor_group.setLayout(editor_layout)
        scroll_layout.addWidget(editor_group)

        # ---- Path & Start ----
        path_group = QGroupBox("文件路径与任务")
        path_layout = QVBoxLayout()

        path_row = QHBoxLayout()
        path_row.addWidget(QLabel("作文文件夹路径"))
        self.path_entry = QLineEdit(self.config["APP"]["ROOT_DIR"])
        path_row.addWidget(self.path_entry, 1)
        btn_browse = QPushButton("浏览")
        btn_browse.setFixedWidth(70)
        btn_browse.clicked.connect(self._browse_folder)
        path_row.addWidget(btn_browse)
        path_layout.addLayout(path_row)

        # Start button (放在路径下方、任务列表前)
        btn_start = QPushButton("🚀 开始处理（自动读取路径下任务并开始）")
        btn_start.setStyleSheet("background-color: #4CAF50; color: white; font-size: 14px; padding: 8px;")
        btn_start.clicked.connect(self._start_processing)
        path_layout.addWidget(btn_start)

        path_group.setLayout(path_layout)
        scroll_layout.addWidget(path_group)

        # Task queue table (任务日志合一)
        self.queue_table = QTableWidget()
        self.queue_table.setColumnCount(10)
        self.queue_table.setHorizontalHeaderLabels(["序号", "学生姓名", "文件路径", "作文名称", "修改前字数", "年级", "线上/线下", "修改后字数", "状态", "实时日志"])
        header_view = self.queue_table.horizontalHeader()
        header_view.setSectionResizeMode(0, QHeaderView.Fixed)
        header_view.setSectionResizeMode(1, QHeaderView.Fixed)
        header_view.setSectionResizeMode(2, QHeaderView.Stretch)
        header_view.setSectionResizeMode(3, QHeaderView.Fixed)
        header_view.setSectionResizeMode(4, QHeaderView.Fixed)
        header_view.setSectionResizeMode(5, QHeaderView.Fixed)
        header_view.setSectionResizeMode(6, QHeaderView.Fixed)
        header_view.setSectionResizeMode(7, QHeaderView.Fixed)
        header_view.setSectionResizeMode(8, QHeaderView.Fixed)
        header_view.setSectionResizeMode(9, QHeaderView.Stretch)
        self.queue_table.setColumnWidth(0, 40)
        self.queue_table.setColumnWidth(1, 100)
        self.queue_table.setColumnWidth(3, 120)
        self.queue_table.setColumnWidth(4, 80)
        self.queue_table.setColumnWidth(5, 80)
        self.queue_table.setColumnWidth(6, 80)
        self.queue_table.setColumnWidth(7, 80)
        self.queue_table.setColumnWidth(8, 70)
        self.queue_table.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.queue_table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.queue_table.verticalHeader().setVisible(False)
        self.queue_table.setMinimumHeight(280)
        self.queue_table.setMaximumHeight(400)
        scroll_layout.addWidget(self.queue_table)

        queue_btns = QHBoxLayout()
        btn_add_task = QPushButton("添加")
        btn_add_task.clicked.connect(self._add_task)
        btn_del_task = QPushButton("删除")
        btn_del_task.clicked.connect(self._remove_task)
        btn_requeue_task = QPushButton("重新加入")
        btn_requeue_task.clicked.connect(self._requeue_selected_tasks)
        btn_load_task = QPushButton("读取")
        btn_load_task.clicked.connect(self._load_tasks)
        btn_refresh_task = QPushButton("刷新队列")
        btn_refresh_task.clicked.connect(self._refresh_queue)
        queue_btns.addWidget(btn_add_task)
        queue_btns.addWidget(btn_del_task)
        queue_btns.addWidget(btn_requeue_task)
        queue_btns.addStretch()
        queue_btns.addWidget(btn_load_task)
        queue_btns.addWidget(btn_refresh_task)
        scroll_layout.addLayout(queue_btns)

        # Log (collapsible)
        self.log_section = CollapsibleSection("运行日志", collapsed=True)
        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        self.log_text.setMaximumHeight(360)
        self.log_section.add_widget(self.log_text)
        scroll_layout.addWidget(self.log_section)

        scroll_layout.addStretch()
        scroll.setWidget(scroll_widget)
        layout.addWidget(scroll)

    def _on_typo_provider_change(self, name):
        p_name = _ensure_provider_exists(self.config, name)
        self.typo_base_entry.setText((self.config.get("LLM", {}).get("PROVIDERS", {}).get(p_name, {}) or {}).get("BASE_URL", ""))

    def _on_editor_provider_change(self, name):
        p_name = _ensure_provider_exists(self.config, name)
        self.editor_base_entry.setText((self.config.get("LLM", {}).get("PROVIDERS", {}).get(p_name, {}) or {}).get("BASE_URL", ""))

    def _add_provider(self, combo):
        from PySide6.QtWidgets import QInputDialog
        name, ok = QInputDialog.getText(self, "新增 AI Provider", "Provider 名称（如 xai / moonshot）:")
        if ok and name.strip():
            p_name = _normalize_provider_exists(self.config, name.strip())
            _ensure_provider_exists(self.config, p_name)
            save_config(self.config)
            names = _provider_name_list(self.config)
            combo.clear()
            combo.addItems(names)
            idx = combo.findText(p_name)
            if idx >= 0:
                combo.setCurrentIndex(idx)

    def _browse_folder(self):
        folder = QFileDialog.getExistingDirectory(self, "选择作文文件夹")
        if folder:
            self.path_entry.setText(folder)
            self._refresh_queue()

    # ---- Task Queue ----
    def _refresh_queue(self):
        with self.queue_lock:
            self.task_queue = [p for p in self.task_queue if os.path.isdir(p)]
        self._render_queue()

    def _render_queue(self):
        old_rows = {}
        for row in range(self.queue_table.rowCount()):
            path_item = self.queue_table.item(row, 2)
            if path_item:
                old_rows[path_item.text()] = {
                    "essay_title": self.queue_table.item(row, 3).text() if self.queue_table.item(row, 3) else "-",
                    "before_count": self.queue_table.item(row, 4).text() if self.queue_table.item(row, 4) else "-",
                    "grade": self.queue_table.item(row, 5).text() if self.queue_table.item(row, 5) else "-",
                    "online_offline": self.queue_table.item(row, 6).text() if self.queue_table.item(row, 6) else "-",
                    "after_count": self.queue_table.item(row, 7).text() if self.queue_table.item(row, 7) else "-",
                    "status": self.queue_table.item(row, 8).text() if self.queue_table.item(row, 8) else "待完成",
                    "log": self.queue_table.item(row, 9).text() if self.queue_table.item(row, 9) else "等待开始...",
                }
        status_colors = {"待完成": "#cfe2ff", "处理中": "#fff3bf", "已完成": "#d4edda", "失败": "#f8d7da"}

        self.queue_table.setRowCount(0)
        with self.queue_lock:
            task_paths = list(self.task_queue)
        for i, task_path in enumerate(task_paths, start=1):
            folder_name = os.path.basename(task_path)
            student = folder_name
            row = self.queue_table.rowCount()
            self.queue_table.insertRow(row)
            self.queue_table.setItem(row, 0, QTableWidgetItem(str(i)))
            self.queue_table.setItem(row, 1, QTableWidgetItem(student))
            self.queue_table.setItem(row, 2, QTableWidgetItem(task_path))
            if task_path in old_rows:
                row_state = old_rows[task_path]
                self.queue_table.setItem(row, 3, QTableWidgetItem(row_state["essay_title"]))
                self.queue_table.setItem(row, 4, QTableWidgetItem(row_state["before_count"]))
                self.queue_table.setItem(row, 5, QTableWidgetItem(row_state["grade"]))
                self.queue_table.setItem(row, 6, QTableWidgetItem(row_state["online_offline"]))
                self.queue_table.setItem(row, 7, QTableWidgetItem(row_state["after_count"]))
                self.queue_table.setItem(row, 8, QTableWidgetItem(row_state["status"]))
                self.queue_table.setItem(row, 9, QTableWidgetItem(row_state["log"]))
                bg_color = QColor(status_colors.get(row_state["status"], "#ffffff"))
            elif task_path in self.completed_tasks:
                self.queue_table.setItem(row, 3, QTableWidgetItem("-"))
                self.queue_table.setItem(row, 4, QTableWidgetItem("-"))
                self.queue_table.setItem(row, 5, QTableWidgetItem("-"))
                self.queue_table.setItem(row, 6, QTableWidgetItem("-"))
                self.queue_table.setItem(row, 7, QTableWidgetItem("-"))
                self.queue_table.setItem(row, 8, QTableWidgetItem("已完成"))
                self.queue_table.setItem(row, 9, QTableWidgetItem("之前已处理完成，跳过"))
                bg_color = QColor("#d4edda")
            else:
                self.queue_table.setItem(row, 3, QTableWidgetItem("-"))
                self.queue_table.setItem(row, 4, QTableWidgetItem("-"))
                self.queue_table.setItem(row, 5, QTableWidgetItem("-"))
                self.queue_table.setItem(row, 6, QTableWidgetItem("-"))
                self.queue_table.setItem(row, 7, QTableWidgetItem("-"))
                self.queue_table.setItem(row, 8, QTableWidgetItem("待完成"))
                self.queue_table.setItem(row, 9, QTableWidgetItem("等待开始..."))
                bg_color = QColor("#cfe2ff")
            for col in range(10):
                item = self.queue_table.item(row, col)
                if item:
                    item.setBackground(bg_color)

    def _add_task(self):
        folder = QFileDialog.getExistingDirectory(self, "选择任务文件夹")
        if not folder:
            return
        folder = os.path.abspath(folder)
        if not os.path.isdir(folder) or not has_images_folder(folder):
            self.log_signal.log_message.emit(f"无效文件夹或无图片：{folder}")
            return
        if folder in self.task_queue:
            self.log_signal.log_message.emit(f"已存在：{folder}")
            return
        with self.queue_lock:
            self.task_queue.append(folder)
            self.finished_tasks.discard(folder)
            self.completed_tasks.discard(folder)
            self.in_progress_tasks.discard(folder)
        self._render_queue()
        self.log_signal.log_message.emit(f"已添加任务：{folder}")

    def _remove_task(self):
        rows = self.queue_table.selectionModel().selectedRows()
        if not rows:
            self.log_signal.log_message.emit("请先选择要删除的队列项")
            return
        for idx in sorted(rows, reverse=True):
            path = self.queue_table.item(idx.row(), 2).text()
            with self.queue_lock:
                if path in self.task_queue:
                    self.task_queue.remove(path)
                # 删除时清除完成标记，重新加入可再处理
                self.completed_tasks.discard(path)
                self.finished_tasks.discard(path)
                self.in_progress_tasks.discard(path)
        self._refresh_queue()

    def _requeue_selected_tasks(self):
        rows = self.queue_table.selectionModel().selectedRows()
        if not rows:
            self.log_signal.log_message.emit("请先选择要重新加入的队列项")
            return

        requeued = 0
        skipped_running = 0
        for idx in rows:
            path_item = self.queue_table.item(idx.row(), 2)
            if not path_item:
                continue
            task_path = path_item.text()
            with self.queue_lock:
                if task_path in self.in_progress_tasks:
                    skipped_running += 1
                    continue
                if task_path not in self.task_queue:
                    self.task_queue.append(task_path)
                self.completed_tasks.discard(task_path)
                self.finished_tasks.discard(task_path)
            self.log_signal.task_status.emit(task_path, "pending", "等待重试", "", "手动重新加入队列", "", "", "", "")
            requeued += 1

        if requeued:
            self.log_signal.log_message.emit(f"已重新加入 {requeued} 个任务")
        if skipped_running:
            self.log_signal.log_message.emit(f"{skipped_running} 个任务正在处理中，已跳过")
        self._refresh_queue()

    def _load_tasks(self):
        folder = self.path_entry.text().strip()
        if not folder or not os.path.isdir(folder):
            self.log_signal.log_message.emit("当前路径无效")
            return
        candidates = scan_folder_for_tasks(folder)
        added = 0
        with self.queue_lock:
            for p in candidates:
                if p not in self.task_queue:
                    self.task_queue.append(p)
                    self.completed_tasks.discard(p)
                    self.finished_tasks.discard(p)
                    self.in_progress_tasks.discard(p)
                    added += 1
        self.log_signal.log_message.emit(f"已读取并加入 {added} 个任务" if added else "没有新任务")
        self._render_queue()

    def _update_task_status(self, task_path: str, status: str, step: str = "", after_count: str = "", log_msg: str = "", grade: str = "", online_offline: str = "", before_count: str = "", essay_title: str = ""):
        """更新任务状态"""
        labels = {"pending": "待完成", "running": "处理中", "done": "已完成", "failed": "失败"}
        colors = {"pending": "#cfe2ff", "running": "#fff3bf", "done": "#d4edda", "failed": "#f8d7da"}
        for row in range(self.queue_table.rowCount()):
            if self.queue_table.item(row, 2) and self.queue_table.item(row, 2).text() == task_path:
                # 实时日志：只显示最新的步骤或日志（替换）
                display_msg = step if step else log_msg
                if display_msg:
                    self.queue_table.item(row, 9).setText(display_msg)
                if essay_title:
                    self.queue_table.item(row, 3).setText(essay_title)
                if before_count:
                    self.queue_table.item(row, 4).setText(before_count)
                if grade:
                    self.queue_table.item(row, 5).setText(grade)
                if online_offline:
                    self.queue_table.item(row, 6).setText(online_offline)
                if after_count:
                    self.queue_table.item(row, 7).setText(after_count)
                if status:
                    self.queue_table.item(row, 8).setText(labels.get(status, status))
                if status:
                    for col in range(10):
                        item = self.queue_table.item(row, col)
                        if item:
                            item.setBackground(QColor(colors.get(status, "#ffffff")))
                # 完成的任务加入标记集合
                if status == "done":
                    with self.queue_lock:
                        self.completed_tasks.add(task_path)
                        self.finished_tasks.add(task_path)
                elif status == "failed":
                    with self.queue_lock:
                        self.finished_tasks.add(task_path)
                elif status == "pending":
                    with self.queue_lock:
                        self.completed_tasks.discard(task_path)
                        self.finished_tasks.discard(task_path)
                        self.in_progress_tasks.discard(task_path)
                break

    # ---- Start Processing ----
    def _start_processing(self):
        # 自动读取路径下的任务加入队列
        folder = self.path_entry.text().strip()
        added = 0
        requeued_tasks = []
        if folder and os.path.isdir(folder):
            candidates = scan_folder_for_tasks(folder)
            with self.queue_lock:
                for task_path in candidates:
                    if task_path not in self.task_queue:
                        self.task_queue.append(task_path)
                        self.completed_tasks.discard(task_path)
                        self.finished_tasks.discard(task_path)
                        self.in_progress_tasks.discard(task_path)
                        added += 1
        with self.queue_lock:
            # 每次点击开始时，只跳过已完成任务；失败/未完成任务重新进入本批次。
            for task_path in self.task_queue:
                if task_path not in self.completed_tasks and task_path not in self.in_progress_tasks:
                    if task_path in self.finished_tasks:
                        requeued_tasks.append(task_path)
                    self.finished_tasks.discard(task_path)
            should_start_worker = not self.is_processing
            if should_start_worker:
                self.is_processing = True
        if added:
            self.log_signal.log_message.emit(f"已自动读取 {added} 个任务")
        for task_path in requeued_tasks:
            self.log_signal.task_status.emit(task_path, "pending", "等待重试", "", "重新加入队列", "", "", "", "")
        self._refresh_queue()

        if not should_start_worker:
            self.log_signal.log_message.emit("当前已有任务在处理，新加入的任务会排队继续处理")
            return

        threading.Thread(target=self._run_processing, daemon=True).start()

    def _run_processing(self):
        cfg = self.config
        cfg.setdefault("OCR", {})
        cfg["OCR"].setdefault("XFYUN", {})
        cfg["OCR"]["PROVIDER"] = "xfyun_handwriting"
        cfg["OCR"]["XFYUN"]["URL"] = self.url_entry.text().strip()
        cfg["OCR"]["XFYUN"]["APPID"] = self.appid_entry.text().strip()
        cfg["OCR"]["XFYUN"]["API_KEY"] = self.apikey_entry.text().strip()
        cfg["OCR"]["XFYUN"].setdefault("LANGUAGE", "cn|en")
        cfg["OCR"]["XFYUN"].setdefault("LOCATION", "false")

        cfg.setdefault("APP", {})
        cfg["APP"]["ROOT_DIR"] = self.path_entry.text().strip()

        cfg.setdefault("LLM", {})
        cfg["LLM"].setdefault("PROVIDERS", {})
        cfg["LLM"].setdefault("TASKS", {})

        typo_provider = _ensure_provider_exists(cfg, self.typo_provider_combo.currentText() or "deepseek")
        cfg["LLM"]["PROVIDERS"].setdefault(typo_provider, {})
        cfg["LLM"]["PROVIDERS"][typo_provider]["API_KEY"] = self.typo_api_key.text().strip()
        cfg["LLM"]["PROVIDERS"][typo_provider]["BASE_URL"] = self.typo_base_entry.text().strip()
        cfg["LLM"]["PROVIDERS"][typo_provider].setdefault("MODEL", "deepseek-chat" if typo_provider == "deepseek" else "gpt-4o-mini")

        cfg["LLM"]["TASKS"].setdefault("typo_fix", {})
        cfg["LLM"]["TASKS"]["typo_fix"]["ENABLED"] = self.use_typo_fix.isChecked()
        cfg["LLM"]["TASKS"]["typo_fix"]["PROVIDER"] = typo_provider
        cfg["LLM"]["TASKS"]["typo_fix"]["PROMPT"] = self.typo_prompt_text.toPlainText().strip()

        editor_provider = _ensure_provider_exists(cfg, self.editor_provider_combo.currentText() or "deepseek")
        cfg["LLM"]["PROVIDERS"].setdefault(editor_provider, {})
        cfg["LLM"]["PROVIDERS"][editor_provider]["API_KEY"] = self.editor_api_key.text().strip()
        cfg["LLM"]["PROVIDERS"][editor_provider]["BASE_URL"] = self.editor_base_entry.text().strip()
        cfg["LLM"]["PROVIDERS"][editor_provider].setdefault("MODEL", "deepseek-chat" if editor_provider == "deepseek" else "gpt-4o-mini")

        count_min = count_max = None
        min_text = self.editor_count_min.text().strip()
        max_text = self.editor_count_max.text().strip()
        if min_text:
            try:
                count_min = int(min_text)
            except ValueError:
                self.log_signal.log_message.emit("目标字数范围最小值必须是整数")
                with self.queue_lock:
                    self.is_processing = False
                return
        if max_text:
            try:
                count_max = int(max_text)
            except ValueError:
                self.log_signal.log_message.emit("目标字数范围最大值必须是整数")
                with self.queue_lock:
                    self.is_processing = False
                return
        if count_min is not None and count_max is not None and count_min > count_max:
            self.log_signal.log_message.emit("目标字数范围最小值不能大于最大值")
            with self.queue_lock:
                self.is_processing = False
            return

        cfg["LLM"]["TASKS"].setdefault("editor", {})
        cfg["LLM"]["TASKS"]["editor"]["ENABLED"] = self.use_editor.isChecked()
        cfg["LLM"]["TASKS"]["editor"]["PROVIDER"] = editor_provider
        cfg["LLM"]["TASKS"]["editor"]["PROMPT"] = self.editor_prompt_text.toPlainText().strip()
        cfg["LLM"]["TASKS"]["editor"]["COUNT_MIN"] = count_min
        cfg["LLM"]["TASKS"]["editor"]["COUNT_MAX"] = count_max

        cfg["OCR"].setdefault("BAIDU_CORRECTION", {})
        cfg["OCR"]["BAIDU_CORRECTION"]["ENABLED"] = self.use_baidu_correction.isChecked()
        cfg["OCR"]["BAIDU_CORRECTION"]["API_KEY"] = self.baidu_api_key_entry.text().strip()
        cfg["OCR"]["BAIDU_CORRECTION"]["SECRET_KEY"] = self.baidu_secret_key_entry.text().strip()

        if not all([cfg.get("OCR", {}).get("XFYUN", {}).get("URL"), cfg.get("OCR", {}).get("XFYUN", {}).get("APPID"), cfg.get("OCR", {}).get("XFYUN", {}).get("API_KEY"), cfg.get("APP", {}).get("ROOT_DIR")]):
            self.log_signal.log_message.emit("请填写完整的 OCR 配置和文件夹路径")
            with self.queue_lock:
                self.is_processing = False
            return

        if not os.path.isdir(cfg["APP"]["ROOT_DIR"]):
            self.log_signal.log_message.emit("文件夹路径无效")
            with self.queue_lock:
                self.is_processing = False
            return

        save_config(cfg)
        tasks_cfg = cfg.get("LLM", {}).get("TASKS", {})
        if tasks_cfg.get("typo_fix", {}).get("ENABLED"):
            self.log_signal.log_message.emit("AI 错别字修正：已启用")
        if cfg.get("OCR", {}).get("BAIDU_CORRECTION", {}).get("ENABLED"):
            self.log_signal.log_message.emit("百度图片矫正：已启用")

        def task_status_cb(folder_path, status, step="", log_msg="", after_count="", grade="", online_offline="", before_count="", essay_title=""):
            self.log_signal.task_status.emit(
                folder_path,
                status,
                step or "",
                str(after_count) if after_count else "",
                log_msg or "",
                grade or "",
                online_offline or "",
                str(before_count) if before_count else "",
                essay_title or "",
            )

        try:
            self.log_signal.log_message.emit(f"开始处理...（并发数：{self.max_parallel_tasks}）")
            from ocr_main import process_folder

            def process_one_task(task_path):
                task_name = os.path.basename(task_path)

                def task_log(msg):
                    self.log_signal.log_message.emit(f"[{task_name}] {msg}")
                    self.log_signal.task_status.emit(task_path, "", "", "", msg or "", "", "", "", "")

                try:
                    self.log_signal.task_status.emit(task_path, "running", "排队启动", "", "开始处理", "", "", "", "")
                    task_log(f"处理: {task_path}")
                    process_folder(
                        task_path,
                        log_callback=task_log,
                        use_typo_fix=bool(tasks_cfg.get("typo_fix", {}).get("ENABLED", False)),
                        use_editor=bool(tasks_cfg.get("editor", {}).get("ENABLED", False)),
                        task_status_callback=task_status_cb,
                    )
                except Exception as exc:
                    self.log_signal.task_status.emit(task_path, "failed", "", "", f"失败: {exc}", "", "", "", "")
                    self.log_signal.log_message.emit(f"[{task_name}] 处理失败：{exc}")

            started_any = False
            futures = {}
            with ThreadPoolExecutor(max_workers=self.max_parallel_tasks) as executor:
                while True:
                    with self.queue_lock:
                        slots = self.max_parallel_tasks - len(futures)
                        pending_tasks = [
                            p for p in self.task_queue
                            if p not in self.finished_tasks and p not in self.in_progress_tasks
                        ]
                        next_tasks = pending_tasks[:max(0, slots)]
                        for task_path in next_tasks:
                            if os.path.isdir(task_path):
                                self.in_progress_tasks.add(task_path)
                            else:
                                self.finished_tasks.add(task_path)

                    for task_path in next_tasks:
                        if not os.path.isdir(task_path):
                            continue
                        started_any = True
                        futures[executor.submit(process_one_task, task_path)] = task_path

                    if not futures:
                        with self.queue_lock:
                            has_more = any(
                                p not in self.finished_tasks and p not in self.in_progress_tasks
                                for p in self.task_queue
                            )
                            if not has_more:
                                self.is_processing = False
                        if not has_more:
                            if not started_any:
                                self.log_signal.log_message.emit("没有待处理的任务（已完成的任务已跳过）")
                            break

                    done_futures, _ = wait(futures.keys(), timeout=0.3, return_when=FIRST_COMPLETED)
                    for future in done_futures:
                        task_path = futures.pop(future)
                        try:
                            future.result()
                        except Exception as exc:
                            self.log_signal.task_status.emit(task_path, "failed", "", "", f"失败: {exc}", "", "", "", "")
                            self.log_signal.log_message.emit(f"[{os.path.basename(task_path)}] 处理失败：{exc}")
                        finally:
                            with self.queue_lock:
                                self.in_progress_tasks.discard(task_path)
                                self.finished_tasks.add(task_path)

            self.log_signal.log_message.emit("全部处理完成")
        except Exception as e:
            self.log_signal.log_message.emit(f"处理失败：{e}")
        finally:
            with self.queue_lock:
                self.in_progress_tasks.clear()
                self.is_processing = False

    # ===================== PAGE 2: AI DOCX =====================
    def _init_page_ai(self):
        layout = QVBoxLayout(self.page_ai)
        layout.setContentsMargins(0, 0, 0, 0)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll_widget = QWidget()
        scroll_layout = QVBoxLayout(scroll_widget)
        scroll_layout.setContentsMargins(4, 4, 4, 4)
        scroll_layout.setSpacing(6)

        # AI Config
        ai_group = QGroupBox("AI 配置")
        ai_layout = QVBoxLayout()

        row1 = QHBoxLayout()
        row1.addWidget(QLabel("处理文件夹"))
        self.ai_path_entry = QLineEdit()
        row1.addWidget(self.ai_path_entry, 1)
        btn_browse = QPushButton("浏览")
        btn_browse.setFixedWidth(70)
        btn_browse.clicked.connect(self._browse_ai_folder)
        row1.addWidget(btn_browse)
        ai_layout.addLayout(row1)

        row2 = QHBoxLayout()
        row2.addWidget(QLabel("AI Provider"))
        self.ai_provider_combo = QComboBox()
        self.ai_provider_combo.addItems(_provider_name_list(self.config))
        ai_provider = _normalize_provider_name((self.config.get("LLM", {}).get("TASKS", {}).get("editor", {}) or {}).get("PROVIDER", "deepseek")) or "deepseek"
        idx = self.ai_provider_combo.findText(ai_provider)
        if idx >= 0:
            self.ai_provider_combo.setCurrentIndex(idx)
        self.ai_provider_combo.currentTextChanged.connect(self._on_ai_provider_change)
        row2.addWidget(self.ai_provider_combo)
        btn_new_ai = QPushButton("+ 新增")
        btn_new_ai.setFixedWidth(68)
        btn_new_ai.clicked.connect(lambda: self._add_provider(self.ai_provider_combo))
        row2.addWidget(btn_new_ai)
        ai_layout.addLayout(row2)

        row3 = QHBoxLayout()
        row3.addWidget(QLabel("AI API Key"))
        self.ai_key_entry = QLineEdit()
        self.ai_key_entry.setEchoMode(QLineEdit.Password)
        self.ai_key_entry.setText((self.config.get("LLM", {}).get("PROVIDERS", {}).get(ai_provider, {}) or {}).get("API_KEY", ""))
        row3.addWidget(self.ai_key_entry, 1)
        self._make_key_toggle(self.ai_key_entry, row3)
        ai_layout.addLayout(row3)

        row4 = QHBoxLayout()
        row4.addWidget(QLabel("Base URL"))
        self.ai_url_entry = QLineEdit((self.config.get("LLM", {}).get("PROVIDERS", {}).get(ai_provider, {}) or {}).get("BASE_URL", ""))
        row4.addWidget(self.ai_url_entry, 1)
        ai_layout.addLayout(row4)

        row5 = QHBoxLayout()
        row5.addWidget(QLabel("AI Prompt"))
        self.ai_prompt_text = QTextEdit()
        self.ai_prompt_text.setMaximumHeight(80)
        self.ai_prompt_text.setPlainText((self.config.get("LLM", {}).get("TASKS", {}).get("editor", {}) or {}).get("PROMPT", "{text}"))
        row5.addWidget(self.ai_prompt_text, 1)
        ai_layout.addLayout(row5)

        row6 = QHBoxLayout()
        row6.addWidget(QLabel("目标字数"))
        self.ai_count_min = QLineEdit(str((self.config.get("LLM", {}).get("TASKS", {}).get("editor", {}) or {}).get("COUNT_MIN") or ""))
        self.ai_count_min.setFixedWidth(100)
        row6.addWidget(self.ai_count_min)
        row6.addWidget(QLabel("-"))
        self.ai_count_max = QLineEdit(str((self.config.get("LLM", {}).get("TASKS", {}).get("editor", {}) or {}).get("COUNT_MAX") or ""))
        self.ai_count_max.setFixedWidth(100)
        row6.addWidget(self.ai_count_max)
        row6.addWidget(QLabel("（空白表示自动）"))
        row6.addStretch()
        ai_layout.addLayout(row6)

        ai_group.setLayout(ai_layout)
        scroll_layout.addWidget(ai_group)

        # Task flow
        flow_group = QGroupBox("处理流程（勾选/取消步骤，拖动上下移动顺序）")
        flow_layout = QVBoxLayout()

        self.task_config = [
            {"id": "6", "name": "6. 转换 DOC -> DOCX", "enabled": True, "order": 0},
            {"id": "1", "name": "1. 清除空格", "enabled": True, "order": 1},
            {"id": "AI", "name": "AI 改作文", "enabled": True, "order": 2},
            {"id": "2", "name": '2. 添加"修改前/后"', "enabled": True, "order": 3},
            {"id": "3", "name": "3. 格式化字体段落", "enabled": True, "order": 4},
            {"id": "5", "name": "5. 修改作者", "enabled": True, "order": 5},
        ]

        self.task_checkboxes = {}
        for task in sorted(self.task_config, key=lambda x: x["order"]):
            cb = QCheckBox(task["name"])
            cb.setChecked(task["enabled"])
            cb.toggled.connect(lambda checked, t=task: t.update({"enabled": checked}))
            flow_layout.addWidget(cb)
            self.task_checkboxes[task["id"]] = cb

        flow_group.setLayout(flow_layout)
        scroll_layout.addWidget(flow_group)

        # Start button
        btn_start_ai = QPushButton("🚀 开始流程")
        btn_start_ai.setStyleSheet("background-color: #2196F3; color: white; font-size: 14px; padding: 8px;")
        btn_start_ai.clicked.connect(self._start_ai_workflow)
        scroll_layout.addWidget(btn_start_ai)

        # Task queue table (status only)
        self.ai_queue_table = QTableWidget()
        self.ai_queue_table.setColumnCount(9)
        self.ai_queue_table.setHorizontalHeaderLabels(["序号", "学生姓名", "文件路径", "作文名称", "修改前字数", "当前步骤", "修改后字数", "状态", "实时日志"])
        ai_header = self.ai_queue_table.horizontalHeader()
        ai_header.setSectionResizeMode(0, QHeaderView.Fixed)
        ai_header.setSectionResizeMode(1, QHeaderView.Fixed)
        ai_header.setSectionResizeMode(2, QHeaderView.Stretch)
        ai_header.setSectionResizeMode(3, QHeaderView.Fixed)
        ai_header.setSectionResizeMode(4, QHeaderView.Fixed)
        ai_header.setSectionResizeMode(5, QHeaderView.Fixed)
        ai_header.setSectionResizeMode(6, QHeaderView.Fixed)
        ai_header.setSectionResizeMode(7, QHeaderView.Fixed)
        ai_header.setSectionResizeMode(8, QHeaderView.Stretch)
        self.ai_queue_table.setColumnWidth(0, 40)
        self.ai_queue_table.setColumnWidth(1, 100)
        self.ai_queue_table.setColumnWidth(3, 120)
        self.ai_queue_table.setColumnWidth(4, 80)
        self.ai_queue_table.setColumnWidth(5, 100)
        self.ai_queue_table.setColumnWidth(6, 80)
        self.ai_queue_table.setColumnWidth(7, 70)
        self.ai_queue_table.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.ai_queue_table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.ai_queue_table.verticalHeader().setVisible(False)
        self.ai_queue_table.setMinimumHeight(240)
        self.ai_queue_table.setMaximumHeight(360)
        scroll_layout.addWidget(self.ai_queue_table)

        # Log (collapsible)
        self.ai_log_section = CollapsibleSection("处理日志", collapsed=True)
        self.ai_log_text = QTextEdit()
        self.ai_log_text.setReadOnly(True)
        self.ai_log_text.setMaximumHeight(200)
        self.ai_log_section.add_widget(self.ai_log_text)
        scroll_layout.addWidget(self.ai_log_section)

        scroll_layout.addStretch()
        scroll.setWidget(scroll_widget)
        layout.addWidget(scroll)

    # ===================== PAGE 3: DOC AI 识别修改 =====================
    def _init_page_doc_ai(self):
        layout = QVBoxLayout(self.page_doc_ai)
        layout.setContentsMargins(0, 0, 0, 0)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll_widget = QWidget()
        scroll_layout = QVBoxLayout(scroll_widget)
        scroll_layout.setContentsMargins(4, 4, 4, 4)
        scroll_layout.setSpacing(6)

        # AI 配置
        ai_group = QGroupBox("AI 配置")
        ai_form = QFormLayout()

        # Provider 选择
        provider_layout = QHBoxLayout()
        self.doc_ai_provider_combo = QComboBox()
        names = _provider_name_list(self.config)
        self.doc_ai_provider_combo.addItems(names)
        self.doc_ai_provider_combo.currentTextChanged.connect(self._on_doc_ai_provider_change)
        provider_layout.addWidget(self.doc_ai_provider_combo, 1)
        ai_form.addRow("Provider:", provider_layout)

        # API Key
        self.doc_ai_key_entry = QLineEdit()
        self.doc_ai_key_entry.setEchoMode(QLineEdit.Password)
        default_provider = _ensure_provider_exists(self.config, "deepseek")
        default_cfg = (self.config.get("LLM", {}).get("PROVIDERS", {}).get(default_provider, {}) or {})
        self.doc_ai_key_entry.setText(default_cfg.get("API_KEY", ""))
        key_row = QHBoxLayout()
        key_row.addWidget(self.doc_ai_key_entry)
        btn_toggle = QPushButton("👁")
        btn_toggle.setFixedWidth(28)
        btn_toggle.setCheckable(True)
        btn_toggle.setStyleSheet("QPushButton { border: none; font-size: 14px; }")
        btn_toggle.toggled.connect(lambda checked: self.doc_ai_key_entry.setEchoMode(QLineEdit.Normal if checked else QLineEdit.Password))
        key_row.addWidget(btn_toggle)
        ai_form.addRow("API Key:", key_row)

        # Base URL
        self.doc_ai_url_entry = QLineEdit()
        self.doc_ai_url_entry.setText(default_cfg.get("BASE_URL", ""))
        ai_form.addRow("Base URL:", self.doc_ai_url_entry)

        # Model
        model_layout = QHBoxLayout()
        self.doc_ai_model_combo = QComboBox()
        self.doc_ai_model_combo.setEditable(True)
        current_model = default_cfg.get("MODEL", "deepseek-chat")
        self.doc_ai_model_combo.setCurrentText(current_model)
        model_layout.addWidget(self.doc_ai_model_combo, 1)
        btn_refresh_models = QPushButton("获取列表")
        btn_refresh_models.setFixedWidth(70)
        btn_refresh_models.clicked.connect(self._refresh_doc_ai_models)
        model_layout.addWidget(btn_refresh_models)
        ai_form.addRow("Model:", model_layout)

        # Prompt
        ai_group.setLayout(ai_form)
        scroll_layout.addWidget(ai_group)

        # 文件夹路径
        path_group = QGroupBox("文件夹路径")
        path_layout = QHBoxLayout()
        self.doc_ai_path_entry = QLineEdit()
        self.doc_ai_path_entry.setPlaceholderText("选择包含docx文件的文件夹...")
        btn_browse = QPushButton("浏览")
        btn_browse.clicked.connect(self._browse_doc_ai_folder)
        path_layout.addWidget(self.doc_ai_path_entry, 1)
        path_layout.addWidget(btn_browse)
        path_group.setLayout(path_layout)
        scroll_layout.addWidget(path_group)

        # 自定义提示词
        prompt_group = QGroupBox("AI 修改作文提示词（自定义）")
        prompt_layout = QVBoxLayout()
        self.doc_ai_prompt_text = QTextEdit()
        self.doc_ai_prompt_text.setMaximumHeight(120)
        self.doc_ai_prompt_text.setPlainText(
            (self.config.get("LLM", {}).get("TASKS", {}).get("editor", {}) or {}).get("PROMPT")
            or DEFAULT_CONFIG["LLM"]["TASKS"]["editor"]["PROMPT"]
        )
        self.doc_ai_prompt_text.setPlaceholderText("输入AI修改作文的提示词，留空使用默认提示词...")
        prompt_layout.addWidget(self.doc_ai_prompt_text)
        prompt_group.setLayout(prompt_layout)
        scroll_layout.addWidget(prompt_group)

        # 并发设置
        parallel_group = QGroupBox("并发设置")
        parallel_layout = QHBoxLayout()
        parallel_layout.addWidget(QLabel("并发数:"))
        self.doc_ai_parallel_spin = QComboBox()
        self.doc_ai_parallel_spin.addItems(["1", "2", "3", "4", "5"])
        self.doc_ai_parallel_spin.setCurrentText("3")
        parallel_layout.addWidget(self.doc_ai_parallel_spin)
        parallel_layout.addStretch()
        parallel_group.setLayout(parallel_layout)
        scroll_layout.addWidget(parallel_group)

        # 字数限制设置
        count_group = QGroupBox("字数限制（修改后字数）")
        count_layout = QHBoxLayout()
        count_layout.addWidget(QLabel("最少字数:"))
        self.doc_ai_count_min = QLineEdit("780")
        self.doc_ai_count_min.setFixedWidth(80)
        count_layout.addWidget(self.doc_ai_count_min)
        count_layout.addWidget(QLabel("-"))
        count_layout.addWidget(QLabel("最多字数:"))
        self.doc_ai_count_max = QLineEdit("930")
        self.doc_ai_count_max.setFixedWidth(80)
        count_layout.addWidget(self.doc_ai_count_max)
        count_layout.addWidget(QLabel("（留空表示不限制）"))
        count_layout.addStretch()
        count_group.setLayout(count_layout)
        scroll_layout.addWidget(count_group)

        # 开始按钮
        btn_start = QPushButton("开始识别")
        btn_start.setStyleSheet("background-color: #4CAF50; color: white; font-size: 14px; padding: 8px;")
        btn_start.clicked.connect(self._start_doc_ai_workflow)
        scroll_layout.addWidget(btn_start)

        # 任务队列表格
        self.doc_ai_queue_table = QTableWidget()
        self.doc_ai_queue_table.setColumnCount(10)
        self.doc_ai_queue_table.setHorizontalHeaderLabels([
            "序号", "文件路径", "作文标题", "作者", "原文字数", "修改后字数", "年级", "线上线下", "是否合格", "状态"
        ])
        header = self.doc_ai_queue_table.horizontalHeader()
        header.setSectionResizeMode(0, QHeaderView.Fixed)
        header.setSectionResizeMode(1, QHeaderView.Stretch)
        header.setSectionResizeMode(2, QHeaderView.Fixed)
        header.setSectionResizeMode(3, QHeaderView.Fixed)
        header.setSectionResizeMode(4, QHeaderView.Fixed)
        header.setSectionResizeMode(5, QHeaderView.Fixed)
        header.setSectionResizeMode(6, QHeaderView.Fixed)
        header.setSectionResizeMode(7, QHeaderView.Fixed)
        header.setSectionResizeMode(8, QHeaderView.Fixed)
        header.setSectionResizeMode(9, QHeaderView.Fixed)
        self.doc_ai_queue_table.setColumnWidth(0, 40)
        self.doc_ai_queue_table.setColumnWidth(2, 150)
        self.doc_ai_queue_table.setColumnWidth(3, 80)
        self.doc_ai_queue_table.setColumnWidth(4, 80)
        self.doc_ai_queue_table.setColumnWidth(5, 80)
        self.doc_ai_queue_table.setColumnWidth(6, 60)
        self.doc_ai_queue_table.setColumnWidth(7, 60)
        self.doc_ai_queue_table.setColumnWidth(8, 70)
        self.doc_ai_queue_table.setColumnWidth(9, 80)
        self.doc_ai_queue_table.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.doc_ai_queue_table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.doc_ai_queue_table.verticalHeader().setVisible(False)
        self.doc_ai_queue_table.setMinimumHeight(240)
        self.doc_ai_queue_table.setMaximumHeight(400)
        scroll_layout.addWidget(self.doc_ai_queue_table)

        # 日志区域
        self.doc_ai_log_section = CollapsibleSection("处理日志", collapsed=True)
        self.doc_ai_log_text = QTextEdit()
        self.doc_ai_log_text.setReadOnly(True)
        self.doc_ai_log_text.setMaximumHeight(200)
        self.doc_ai_log_section.add_widget(self.doc_ai_log_text)
        scroll_layout.addWidget(self.doc_ai_log_section)

        scroll_layout.addStretch()
        scroll.setWidget(scroll_widget)
        layout.addWidget(scroll)

    def _on_doc_ai_provider_change(self, name):
        p_name = _ensure_provider_exists(self.config, name)
        provider_cfg = (self.config.get("LLM", {}).get("PROVIDERS", {}).get(p_name, {}) or {})
        self.doc_ai_key_entry.setText(provider_cfg.get("API_KEY", ""))
        self.doc_ai_url_entry.setText(provider_cfg.get("BASE_URL", ""))
        self.doc_ai_model_combo.setCurrentText(provider_cfg.get("MODEL", ""))

    def _refresh_doc_ai_models(self):
        """Fetch model list from the current provider's API and populate the combobox."""
        api_key = self.doc_ai_key_entry.text().strip()
        base_url = self.doc_ai_url_entry.text().strip()
        if not api_key:
            QMessageBox.information(self, "提示", "请先填写 API Key")
            return
        try:
            models = fetch_models(api_key, base_url)
            if models:
                self.doc_ai_model_combo.clear()
                self.doc_ai_model_combo.addItems(models)
                current = self.doc_ai_model_combo.currentText()
                if not current:
                    self.doc_ai_model_combo.setCurrentIndex(0)
            else:
                QMessageBox.information(self, "提示", "未获取到模型列表，请检查 API Key 和 Base URL")
        except Exception as e:
            QMessageBox.warning(self, "获取失败", f"获取模型列表失败：{str(e)}")

    def _browse_doc_ai_folder(self):
        folder = QFileDialog.getExistingDirectory(self, "选择文件夹")
        if folder:
            self.doc_ai_path_entry.setText(folder)

    def _scan_docx_files(self, folder):
        """递归扫描文件夹中的docx文件，排除'修改后'文件夹"""
        docx_files = []
        for root, dirs, files in os.walk(folder):
            # 排除"修改后"文件夹
            dirs[:] = [d for d in dirs if d != "修改后"]
            for file in files:
                if file.lower().endswith(".docx") and not file.startswith("~$"):
                    docx_files.append(os.path.join(root, file))
        return docx_files

    def _start_doc_ai_workflow(self):
        threading.Thread(target=self._run_doc_ai_workflow, daemon=True).start()

    def _process_single_docx(self, docx_path, client, model, custom_prompt, count_min=None, count_max=None, max_retries=3):
        """处理单个docx文件（支持重试）"""
        for attempt in range(max_retries):
            try:
                # 读取docx内容
                try:
                    doc = Document(docx_path)
                except Exception as e:
                    if "core.xml" in str(e) or "archive" in str(e).lower():
                        self.log_signal.doc_ai_log_message.emit(f"  {os.path.basename(docx_path)}: docx文件格式不标准，尝试修复...")
                        # 尝试修复：复制文件并重新保存
                        import tempfile
                        temp_path = tempfile.mktemp(suffix=".docx")
                        shutil.copy2(docx_path, temp_path)
                        try:
                            doc = Document(temp_path)
                            # 保存修复后的文件
                            doc.save(docx_path)
                        except Exception:
                            os.remove(temp_path)
                            raise
                        finally:
                            if os.path.exists(temp_path):
                                os.remove(temp_path)
                    else:
                        raise
                
                full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

                if not full_text.strip():
                    self.log_signal.doc_ai_log_message.emit(f"  {os.path.basename(docx_path)}: 空文档，跳过")
                    self.log_signal.doc_ai_task_status.emit(docx_path, "跳过", "空文档")
                    return

                # 创建"修改后"文件夹
                parent_dir = os.path.dirname(docx_path)
                modified_dir = os.path.join(parent_dir, "修改后")
                os.makedirs(modified_dir, exist_ok=True)

                # AI识别提示词（包含文件路径用于参考）
                identify_prompt = f"""请分析以下文档内容，识别并提取以下信息：
1. 作文标题：从文章中识别的标题，去掉"题目"、"标题"等前缀。也可以从文件路径中参考文件名获取标题。
2. 作者：从文章中识别的作者姓名，去掉"——"前缀。也可以从文件路径中参考文件名获取作者。
3. 原文字数：文章原始字数（包含标点，不含空格）
4. 修改后字数：文章修改后字数（包含标点，不含空格），如果文章没有修改后内容，则填写与原文字数相同
5. 年级：从文章中识别的年级，如：三年级、四年级等，如无法识别则填空字符串""
6. 第几次：从文章或文件路径中识别是第几次作文，如无法识别则填空字符串""
7. 线上或线下：从文章中识别的线上或线下，如无法识别则填空字符串""
8. 有修改后内容：判断文章是否包含修改后的内容，true或false

文档路径：{docx_path}
文件名：{os.path.basename(docx_path)}
所在文件夹：{parent_dir}

文档内容：
{full_text}

请以JSON格式返回，格式如下：
{{
  "作文标题": "识别的标题（必填）",
  "作者": "识别的作者（必填）",
  "原文字数": "原始字数",
  "修改后字数": "修改后字数",
  "年级": "识别的年级或空字符串",
  "第几次": "识别的第几次或空字符串",
  "线上或线下": "识别的线上或线下或空字符串",
  "有修改后内容": true或false,
  "修改前正文": "提取的修改前纯正文内容（不含标题、作者、班级、姓名、题目等元数据，只保留正文段落）",
  "修改后正文": "提取的修改后纯正文内容（如果没有修改后内容则为空字符串）"
}}

重要提示：
1. 修改前正文和修改后正文必须是纯正文内容，不能包含"标题"、"作者"、"班级"、"姓名"、"题目"、"修改前"、"修改后"等元数据
2. 如果文档中有"班级："、"姓名："、"题目："等行，请忽略这些行，只提取正文段落
3. 正文段落通常以"　　"（两个全角空格）开头或自然段落"""

                # 调用AI识别
                self.log_signal.doc_ai_task_status.emit(docx_path, "处理中", "正在调用AI识别...")
                response = client.chat.completions.create(
                    model=model,
                    messages=[
                        {"role": "system", "content": "你是一名专业的文档分析助手，擅长从中文作文中提取信息。请务必返回有效的JSON格式。"},
                        {"role": "user", "content": identify_prompt}
                    ],
                    temperature=0.1,
                    stream=False
                )

                result_text = response.choices[0].message.content.strip()

                # 解析JSON
                try:
                    # 提取JSON内容（可能被```json```包裹）
                    json_match = re.search(r'```json\s*(.*?)\s*```', result_text, re.DOTALL)
                    if json_match:
                        json_str = json_match.group(1)
                    else:
                        json_str = result_text

                    result_data = json.loads(json_str)

                    # 更新表格
                    title = result_data.get("作文标题", "") or ""
                    author = result_data.get("作者", "") or ""
                    original_count = result_data.get("原文字数", "未知")
                    modified_count = result_data.get("修改后字数", "未知")
                    grade = result_data.get("年级", "") or ""
                    times = result_data.get("第几次", "") or ""
                    online_offline = result_data.get("线上或线下", "") or ""
                    has_modified = result_data.get("有修改后内容", False)
                    before_text = result_data.get("修改前正文", "") or ""  # AI提取的修改前纯正文
                    after_text = result_data.get("修改后正文", "") or ""   # AI提取的修改后纯正文

                    # 如果AI未提取到修改前正文，使用原文
                    if not before_text.strip():
                        before_text = full_text

                    # 如果标题或作者为空，尝试从文件名解析
                    if not title or not author:
                        file_stem = os.path.splitext(os.path.basename(docx_path))[0]
                        # 尝试解析文件名格式：作文标题——作者 或 作文标题——作者年级第几次线上线下
                        name_match = re.match(r'^(.+?)——(.+?)(年级|第|$)', file_stem)
                        if name_match:
                            if not title:
                                title = name_match.group(1).strip()
                            if not author:
                                author = name_match.group(2).strip()
                        elif not title:
                            title = file_stem

                    # 确保标题和作者不为空
                    if not title:
                        title = "未知标题"
                    if not author:
                        author = "未知作者"

                    # 生成新文件名：改 [作文标题]——[作者][年级][第几次][线上线下].docx
                    name_parts = [f"改 {title}——{author}"]
                    if grade:
                        name_parts.append(grade)
                    if times:
                        name_parts.append(times)
                    if online_offline:
                        name_parts.append(online_offline)
                    new_filename = "".join(name_parts) + ".docx"
                    new_path = os.path.join(modified_dir, new_filename)

                    # 判断修改后字数是否合格（780-930字）
                    is_qualified = "未知"
                    try:
                        modified_count_int = int(modified_count)
                        if 780 <= modified_count_int <= 930:
                            is_qualified = "合格"
                        else:
                            is_qualified = "不合格"
                    except (ValueError, TypeError):
                        is_qualified = "未知"

                    # 处理文档：如果没有修改后内容，用AI修改
                    if not has_modified:
                        self.log_signal.doc_ai_log_message.emit(f"  {os.path.basename(docx_path)}: 无修改后内容，AI修改中...")
                        self.log_signal.doc_ai_task_status.emit(docx_path, "处理中", "AI修改作文中...")
                        
                        # 构建字数限制说明
                        count_requirement = ""
                        if count_min and count_max:
                            count_requirement = f"\n6. 修改后的正文字数必须控制在 {count_min} 到 {count_max} 字之间（包含标点，不含空格）"
                        elif count_min:
                            count_requirement = f"\n6. 修改后的正文字数必须不少于 {count_min} 字（包含标点，不含空格）"
                        elif count_max:
                            count_requirement = f"\n6. 修改后的正文字数必须不超过 {count_max} 字（包含标点，不含空格）"

                        # 使用自定义提示词，明确要求只输出纯正文
                        if "{text}" in custom_prompt:
                            modify_prompt = custom_prompt.format(text=before_text)
                        else:
                            modify_prompt = custom_prompt + f"""

文档路径：{docx_path}

原始文章正文：
{before_text}

重要要求：
1. 只输出修改后的纯正文内容
2. 不要包含标题、作者、班级、姓名、题目等元数据
3. 不要包含"修改前"、"修改后"等标签
4. 只输出正文段落，每段保持原有格式
5. 不要添加任何解释{count_requirement}"""

                        # 调用AI修改
                        modify_response = client.chat.completions.create(
                            model=model,
                            messages=[
                                {"role": "system", "content": "你是一名严谨的中文校对助手。只输出纯正文内容，不要包含任何元数据或标签。"},
                                {"role": "user", "content": modify_prompt}
                            ],
                            temperature=0.1,
                            stream=False
                        )
                        after_text = modify_response.choices[0].message.content.strip()
                    else:
                        # 原文已有修改后内容，使用AI提取的修改后正文
                        if after_text.strip():
                            # 已经从AI识别中获取到修改后正文
                            pass
                        else:
                            # 尝试从原文中分离修改前和修改后内容
                            lines = full_text.split("\n")
                            after_lines = []
                            is_before = True
                            
                            for line in lines:
                                if "修改前" in line:
                                    is_before = True
                                    continue
                                elif "修改后" in line:
                                    is_before = False
                                    continue
                                
                                if not is_before:
                                    after_lines.append(line)
                            
                            if after_lines:
                                after_text = "\n".join(after_lines)
                            else:
                                # 没有找到修改后内容，用AI修改
                                self.log_signal.doc_ai_log_message.emit(f"  {os.path.basename(docx_path)}: 解析失败，AI修改中...")
                                
                                # 构建字数限制说明
                                count_requirement = ""
                                if count_min and count_max:
                                    count_requirement = f"\n6. 修改后的正文字数必须控制在 {count_min} 到 {count_max} 字之间（包含标点，不含空格）"
                                elif count_min:
                                    count_requirement = f"\n6. 修改后的正文字数必须不少于 {count_min} 字（包含标点，不含空格）"
                                elif count_max:
                                    count_requirement = f"\n6. 修改后的正文字数必须不超过 {count_max} 字（包含标点，不含空格）"

                                # 使用自定义提示词
                                if "{text}" in custom_prompt:
                                    modify_prompt = custom_prompt.format(text=before_text)
                                else:
                                    modify_prompt = custom_prompt + f"""

文档路径：{docx_path}

原始文章正文：
{before_text}

重要要求：
1. 只输出修改后的纯正文内容
2. 不要包含标题、作者、班级、姓名、题目等元数据
3. 不要包含"修改前"、"修改后"等标签
4. 只输出正文段落
5. 不要添加任何解释{count_requirement}"""

                                modify_response = client.chat.completions.create(
                                    model=model,
                                    messages=[
                                        {"role": "system", "content": "你是一名严谨的中文校对助手。只输出纯正文内容，不要包含任何元数据或标签。"},
                                        {"role": "user", "content": modify_prompt}
                                    ],
                                    temperature=0.1,
                                    stream=False
                                )
                                after_text = modify_response.choices[0].message.content.strip()

                    # 创建新文档，格式为：修改前（分页符）修改后
                    # 格式与"图片转作文"功能一致
                    new_doc = Document()
                    
                    # 设置默认字体
                    style = new_doc.styles["Normal"]
                    style.font.name = "宋体"
                    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    style.font.size = Pt(12)

                    # 获取标题和作者
                    title_display = title if title and title != "未知" else os.path.splitext(os.path.basename(docx_path))[0]
                    author_display = author if author and author != "未知" else ""

                    # 添加"修改前："标签
                    p_before_label = new_doc.add_paragraph("修改前：")
                    p_before_label.paragraph_format.first_line_indent = Cm(0.74)
                    p_before_label.paragraph_format.space_before = Pt(0)
                    p_before_label.paragraph_format.space_after = Pt(0)
                    p_before_label.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
                    p_before_label.paragraph_format.line_spacing = Pt(12)

                    # 添加作文标题（居中）
                    p_title = new_doc.add_paragraph(title_display)
                    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_title.paragraph_format.space_before = Pt(0)
                    p_title.paragraph_format.space_after = Pt(0)
                    p_title.paragraph_format.line_spacing = Pt(12)

                    # 添加作者姓名（居中，格式为 "——作者"）
                    if author_display:
                        p_author = new_doc.add_paragraph(f"——{author_display}")
                        p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_author.paragraph_format.space_before = Pt(0)
                        p_author.paragraph_format.space_after = Pt(0)
                        p_author.paragraph_format.line_spacing = Pt(12)

                    # 添加修改前内容（使用AI提取的纯正文）
                    for line in before_text.split("\n"):
                        if line.strip():
                            # 跳过元数据行
                            line_content = line.strip()
                            if any(keyword in line_content for keyword in ["班级：", "姓名：", "题目：", "修改前：", "修改后："]):
                                continue
                            p = new_doc.add_paragraph(line_content)
                            p.paragraph_format.first_line_indent = Cm(0.74)
                            p.paragraph_format.space_before = Pt(0)
                            p.paragraph_format.space_after = Pt(0)
                            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
                            p.paragraph_format.line_spacing = Pt(12)

                    # 添加分页符
                    new_doc.add_page_break()

                    # 添加"修改后："标签
                    p_after_label = new_doc.add_paragraph("修改后：")
                    p_after_label.paragraph_format.first_line_indent = Cm(0.74)
                    p_after_label.paragraph_format.space_before = Pt(0)
                    p_after_label.paragraph_format.space_after = Pt(0)
                    p_after_label.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
                    p_after_label.paragraph_format.line_spacing = Pt(12)

                    # 添加作文标题（居中）
                    p_title2 = new_doc.add_paragraph(title_display)
                    p_title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_title2.paragraph_format.space_before = Pt(0)
                    p_title2.paragraph_format.space_after = Pt(0)
                    p_title2.paragraph_format.line_spacing = Pt(12)

                    # 添加作者姓名（居中，格式为 "——作者"）
                    if author_display:
                        p_author2 = new_doc.add_paragraph(f"——{author_display}")
                        p_author2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_author2.paragraph_format.space_before = Pt(0)
                        p_author2.paragraph_format.space_after = Pt(0)
                        p_author2.paragraph_format.line_spacing = Pt(12)

                    # 添加修改后内容（使用AI提取的纯正文）
                    for line in after_text.split("\n"):
                        if line.strip():
                            # 跳过元数据行
                            line_content = line.strip()
                            if any(keyword in line_content for keyword in ["班级：", "姓名：", "题目：", "修改前：", "修改后："]):
                                continue
                            p = new_doc.add_paragraph(line_content)
                            p.paragraph_format.first_line_indent = Cm(0.74)
                            p.paragraph_format.space_before = Pt(0)
                            p.paragraph_format.space_after = Pt(0)
                            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
                            p.paragraph_format.line_spacing = Pt(12)

                    # 保存新文档
                    new_doc.save(new_path)

                    self.log_signal.doc_ai_task_status.emit(
                        docx_path, "完成",
                        f"标题:{title} 作者:{author} 字数:{original_count} 年级:{grade} 修改后:{modified_count}字 [{is_qualified}]"
                    )
                    self.log_signal.doc_ai_log_message.emit(f"  {os.path.basename(docx_path)}: 已保存到 {new_path}")

                    # 更新表格显示
                    QTimer.singleShot(0, lambda p=docx_path, t=title, a=author, 
                                     oc=original_count, mc=modified_count, 
                                     g=grade, oo=online_offline, iq=is_qualified: 
                                     self._update_doc_ai_table(p, t, a, oc, mc, g, oo, iq))
                    return  # 成功处理，退出重试循环

                except json.JSONDecodeError as e:
                    if attempt < max_retries - 1:
                        self.log_signal.doc_ai_log_message.emit(f"  {os.path.basename(docx_path)}: JSON解析失败，重试中...")
                        continue
                    self.log_signal.doc_ai_log_message.emit(f"  {os.path.basename(docx_path)}: JSON解析失败 - {e}")
                    self.log_signal.doc_ai_task_status.emit(docx_path, "失败", f"JSON解析失败: {e}")

            except Exception as e:
                if attempt < max_retries - 1:
                    self.log_signal.doc_ai_log_message.emit(f"  {os.path.basename(docx_path)}: 处理失败，重试中...")
                    continue
                self.log_signal.doc_ai_log_message.emit(f"  {os.path.basename(docx_path)}: {e}")
                self.log_signal.doc_ai_task_status.emit(docx_path, "失败", str(e))
                if _is_bad_marshal_error(e):
                    _clear_project_bytecode_caches()
                    self.log_signal.doc_ai_log_message.emit(_format_exception_for_log(e))

    def _run_doc_ai_workflow(self):
        import shutil
        
        folder = self.doc_ai_path_entry.text().strip()
        api_key = self.doc_ai_key_entry.text().strip()
        base_url = self.doc_ai_url_entry.text().strip()
        model = self.doc_ai_model_combo.currentText().strip() or "deepseek-chat"
        max_parallel = int(self.doc_ai_parallel_spin.currentText())

        if not folder or not api_key:
            self.log_signal.doc_ai_log_message.emit("请填写文件夹路径和 API Key")
            return
        if not os.path.isdir(folder):
            self.log_signal.doc_ai_log_message.emit("文件夹路径无效")
            return

        # 保存配置
        selected_provider = _ensure_provider_exists(self.config, self.doc_ai_provider_combo.currentText() or "deepseek")
        cfg = self.config
        cfg.setdefault("LLM", {})
        cfg["LLM"].setdefault("PROVIDERS", {})
        cfg["LLM"]["PROVIDERS"].setdefault(selected_provider, {})
        cfg["LLM"]["PROVIDERS"][selected_provider]["API_KEY"] = api_key
        cfg["LLM"]["PROVIDERS"][selected_provider]["BASE_URL"] = base_url
        cfg["LLM"]["PROVIDERS"][selected_provider]["MODEL"] = model
        save_config(cfg)

        # 获取自定义提示词
        custom_prompt = self.doc_ai_prompt_text.toPlainText().strip()
        if not custom_prompt:
            custom_prompt = "下面是一篇中文文章，请你【只修改错别字和明显的识别错误】。\n要求：1. 不改变原意 2. 不润色文风 3. 不增删内容 4. 保持原有段落结构 5. 只输出修改后的完整文章正文\n"

        # 获取字数限制
        count_min = None
        count_max = None
        min_text = self.doc_ai_count_min.text().strip()
        max_text = self.doc_ai_count_max.text().strip()
        if min_text:
            try:
                count_min = int(min_text)
            except ValueError:
                self.log_signal.doc_ai_log_message.emit("最少字数必须是整数")
                return
        if max_text:
            try:
                count_max = int(max_text)
            except ValueError:
                self.log_signal.doc_ai_log_message.emit("最多字数必须是整数")
                return
        if count_min and count_max and count_min > count_max:
            self.log_signal.doc_ai_log_message.emit("最少字数不能大于最多字数")
            return

        # 扫描docx文件
        self.log_signal.doc_ai_log_message.emit("开始扫描文件夹...")
        docx_files = self._scan_docx_files(folder)

        if not docx_files:
            self.log_signal.doc_ai_log_message.emit("未找到docx文件")
            return

        self.log_signal.doc_ai_log_message.emit(f"找到 {len(docx_files)} 个docx文件，开始处理（并发数：{max_parallel}）")
        self.log_signal.doc_ai_tasks_loaded.emit(docx_files)

        # 创建AI客户端
        try:
            client = OpenAI(api_key=api_key, base_url=base_url)
        except Exception as e:
            self.log_signal.doc_ai_log_message.emit(f"创建AI客户端失败: {e}")
            return

        # 使用线程池处理文件
        with ThreadPoolExecutor(max_workers=max_parallel) as executor:
            futures = []
            for docx_path in docx_files:
                future = executor.submit(self._process_single_docx, docx_path, client, model, custom_prompt, count_min, count_max)
                futures.append((future, docx_path))

            # 等待所有任务完成
            for future, docx_path in futures:
                try:
                    future.result()
                except Exception as e:
                    self.log_signal.doc_ai_log_message.emit(f"  {os.path.basename(docx_path)}: 线程异常 - {e}")
                    self.log_signal.doc_ai_task_status.emit(docx_path, "失败", str(e))

        self.log_signal.doc_ai_log_message.emit("处理完成！")

    def _update_doc_ai_table(self, file_path, title, author, original_count, modified_count, grade, online_offline, is_qualified):
        """更新表格中指定文件的信息"""
        for row in range(self.doc_ai_queue_table.rowCount()):
            path_item = self.doc_ai_queue_table.item(row, 1)
            if path_item and path_item.text() == file_path:
                self.doc_ai_queue_table.setItem(row, 2, QTableWidgetItem(title))
                self.doc_ai_queue_table.setItem(row, 3, QTableWidgetItem(author))
                self.doc_ai_queue_table.setItem(row, 4, QTableWidgetItem(str(original_count)))
                self.doc_ai_queue_table.setItem(row, 5, QTableWidgetItem(str(modified_count)))
                self.doc_ai_queue_table.setItem(row, 6, QTableWidgetItem(grade))
                self.doc_ai_queue_table.setItem(row, 7, QTableWidgetItem(online_offline))
                # 设置"是否合格"列，合格显示绿色，不合格显示红色
                qualified_item = QTableWidgetItem(is_qualified)
                if is_qualified == "合格":
                    qualified_item.setBackground(QColor("#d4edda"))  # 浅绿色
                elif is_qualified == "不合格":
                    qualified_item.setBackground(QColor("#f8d7da"))  # 浅红色
                self.doc_ai_queue_table.setItem(row, 8, qualified_item)
                break

    def _render_doc_ai_queue(self, task_paths):
        """渲染任务队列（替换所有任务）"""
        self.doc_ai_queue_table.setRowCount(0)
        for i, task_path in enumerate(task_paths, start=1):
            row = self.doc_ai_queue_table.rowCount()
            self.doc_ai_queue_table.insertRow(row)
            self.doc_ai_queue_table.setItem(row, 0, QTableWidgetItem(str(i)))
            self.doc_ai_queue_table.setItem(row, 1, QTableWidgetItem(task_path))
            self.doc_ai_queue_table.setItem(row, 2, QTableWidgetItem("-"))
            self.doc_ai_queue_table.setItem(row, 3, QTableWidgetItem("-"))
            self.doc_ai_queue_table.setItem(row, 4, QTableWidgetItem("-"))
            self.doc_ai_queue_table.setItem(row, 5, QTableWidgetItem("-"))
            self.doc_ai_queue_table.setItem(row, 6, QTableWidgetItem("-"))
            self.doc_ai_queue_table.setItem(row, 7, QTableWidgetItem("-"))
            self.doc_ai_queue_table.setItem(row, 8, QTableWidgetItem("-"))
            self.doc_ai_queue_table.setItem(row, 9, QTableWidgetItem("待处理"))
            for col in range(10):
                item = self.doc_ai_queue_table.item(row, col)
                if item:
                    item.setBackground(QColor("#cfe2ff"))

    def _add_doc_ai_tasks(self, task_paths):
        """添加任务到队列（不替换现有任务）"""
        existing_paths = set()
        for row in range(self.doc_ai_queue_table.rowCount()):
            path_item = self.doc_ai_queue_table.item(row, 1)
            if path_item:
                existing_paths.add(path_item.text())
        
        added_count = 0
        for task_path in task_paths:
            if task_path not in existing_paths:
                row = self.doc_ai_queue_table.rowCount()
                self.doc_ai_queue_table.insertRow(row)
                self.doc_ai_queue_table.setItem(row, 0, QTableWidgetItem(str(row + 1)))
                self.doc_ai_queue_table.setItem(row, 1, QTableWidgetItem(task_path))
                self.doc_ai_queue_table.setItem(row, 2, QTableWidgetItem("-"))
                self.doc_ai_queue_table.setItem(row, 3, QTableWidgetItem("-"))
                self.doc_ai_queue_table.setItem(row, 4, QTableWidgetItem("-"))
                self.doc_ai_queue_table.setItem(row, 5, QTableWidgetItem("-"))
                self.doc_ai_queue_table.setItem(row, 6, QTableWidgetItem("-"))
                self.doc_ai_queue_table.setItem(row, 7, QTableWidgetItem("-"))
                self.doc_ai_queue_table.setItem(row, 8, QTableWidgetItem("-"))
                self.doc_ai_queue_table.setItem(row, 9, QTableWidgetItem("待处理"))
                for col in range(10):
                    item = self.doc_ai_queue_table.item(row, col)
                    if item:
                        item.setBackground(QColor("#cfe2ff"))
                added_count += 1
        
        # 更新序号
        for row in range(self.doc_ai_queue_table.rowCount()):
            self.doc_ai_queue_table.setItem(row, 0, QTableWidgetItem(str(row + 1)))
        
        return added_count

    def _update_doc_ai_task_status(self, task_path, status, log_msg=""):
        colors = {"待处理": "#cfe2ff", "处理中": "#fff3bf", "完成": "#d4edda", "失败": "#f8d7da", "跳过": "#e2e3e5"}
        for row in range(self.doc_ai_queue_table.rowCount()):
            path_item = self.doc_ai_queue_table.item(row, 1)
            if path_item and path_item.text() == task_path:
                self.doc_ai_queue_table.setItem(row, 9, QTableWidgetItem(status))
                bg = QColor(colors.get(status, "#ffffff"))
                for col in range(10):
                    item = self.doc_ai_queue_table.item(row, col)
                    if item:
                        item.setBackground(bg)
                break

    def _count_docx_file_chars(self, docx_path: str) -> str:
        try:
            if os.path.isfile(docx_path) and docx_path.lower().endswith(".docx"):
                doc = Document(docx_path)
                total = sum(len(p.text.strip()) for p in doc.paragraphs if p.text.strip())
                return str(total)
        except Exception:
            pass
        return ""

    def _ai_output_docx_path(self, task_path: str) -> str:
        root, ext = os.path.splitext(task_path)
        if ext.lower() == ".docx":
            return task_path
        converted = root + ".docx"
        return converted if os.path.isfile(converted) else task_path

    def _render_ai_queue(self, task_paths):
        self.ai_queue_table.setRowCount(0)
        for i, task_path in enumerate(task_paths, start=1):
            base = os.path.basename(task_path)
            display_name = base[2:] if base.startswith("改 ") else base
            stem, ext = os.path.splitext(display_name)
            student, essay = infer_student_and_essay(stem)
            before_count = self._count_docx_file_chars(task_path)
            row = self.ai_queue_table.rowCount()
            self.ai_queue_table.insertRow(row)
            self.ai_queue_table.setItem(row, 0, QTableWidgetItem(str(i)))
            self.ai_queue_table.setItem(row, 1, QTableWidgetItem(student))
            self.ai_queue_table.setItem(row, 2, QTableWidgetItem(task_path))
            self.ai_queue_table.setItem(row, 3, QTableWidgetItem(essay))
            self.ai_queue_table.setItem(row, 4, QTableWidgetItem(before_count))
            self.ai_queue_table.setItem(row, 5, QTableWidgetItem("-"))
            self.ai_queue_table.setItem(row, 6, QTableWidgetItem("-"))
            self.ai_queue_table.setItem(row, 7, QTableWidgetItem("待完成"))
            self.ai_queue_table.setItem(row, 8, QTableWidgetItem("等待开始..."))
            for col in range(9):
                item = self.ai_queue_table.item(row, col)
                if item:
                    item.setBackground(QColor("#cfe2ff"))

    def _update_ai_task_status(self, task_path: str, status: str, step: str = "", after_count: str = "", log_msg: str = ""):
        labels = {"pending": "待完成", "running": "处理中", "done": "已完成", "failed": "失败"}
        colors = {"pending": "#cfe2ff", "running": "#fff3bf", "done": "#d4edda", "failed": "#f8d7da"}
        for row in range(self.ai_queue_table.rowCount()):
            path_item = self.ai_queue_table.item(row, 2)
            if path_item and path_item.text() == task_path:
                if step:
                    self.ai_queue_table.item(row, 5).setText(step)
                if after_count:
                    self.ai_queue_table.item(row, 6).setText(after_count)
                if status:
                    self.ai_queue_table.item(row, 7).setText(labels.get(status, status))
                if log_msg:
                    old_log = self.ai_queue_table.item(row, 8).text()
                    if old_log == "等待开始...":
                        self.ai_queue_table.item(row, 8).setText(log_msg)
                    else:
                        self.ai_queue_table.item(row, 8).setText(old_log + "\n" + log_msg)
                    self.ai_queue_table.scrollToItem(self.ai_queue_table.item(row, 8))
                if status:
                    bg = QColor(colors.get(status, "#ffffff"))
                    for col in range(9):
                        item = self.ai_queue_table.item(row, col)
                        if item:
                            item.setBackground(bg)
                break

    def _browse_ai_folder(self):
        folder = QFileDialog.getExistingDirectory(self, "选择处理文件夹")
        if folder:
            self.ai_path_entry.setText(folder)

    def _on_ai_provider_change(self, name):
        p_name = _ensure_provider_exists(self.config, name)
        provider_cfg = (self.config.get("LLM", {}).get("PROVIDERS", {}).get(p_name, {}) or {})
        self.ai_key_entry.setText(provider_cfg.get("API_KEY", ""))
        self.ai_url_entry.setText(provider_cfg.get("BASE_URL", ""))

    def _start_ai_workflow(self):
        threading.Thread(target=self._run_ai_workflow, daemon=True).start()

    def _run_ai_workflow(self):
        _clear_project_bytecode_caches()
        folder = self.ai_path_entry.text().strip()
        selected_provider = _ensure_provider_exists(self.config, self.ai_provider_combo.currentText() or "deepseek")
        api_key = self.ai_key_entry.text().strip()
        base_url = self.ai_url_entry.text().strip()
        prompt = self.ai_prompt_text.toPlainText().strip() or None

        min_text = self.ai_count_min.text().strip()
        max_text = self.ai_count_max.text().strip()
        count_min = count_max = None
        if min_text:
            try:
                count_min = int(min_text)
            except ValueError:
                self.log_signal.log_message.emit("目标字数范围最小值必须是整数")
                return
        if max_text:
            try:
                count_max = int(max_text)
            except ValueError:
                self.log_signal.log_message.emit("目标字数范围最大值必须是整数")
                return

        if not folder or not api_key:
            self.log_signal.ai_log_message.emit("请填写文件夹路径和 API Key")
            return
        if not os.path.isdir(folder):
            self.log_signal.ai_log_message.emit("文件夹路径无效")
            return

        cfg = self.config
        cfg.setdefault("LLM", {})
        cfg["LLM"].setdefault("PROVIDERS", {})
        cfg["LLM"].setdefault("TASKS", {})
        cfg["LLM"]["PROVIDERS"].setdefault(selected_provider, {})
        cfg["LLM"]["PROVIDERS"][selected_provider]["API_KEY"] = api_key
        cfg["LLM"]["PROVIDERS"][selected_provider]["BASE_URL"] = base_url
        cfg["LLM"]["PROVIDERS"][selected_provider].setdefault("MODEL", "deepseek-chat" if selected_provider == "deepseek" else "gpt-4o-mini")
        cfg["LLM"]["TASKS"].setdefault("editor", {})
        cfg["LLM"]["TASKS"]["editor"]["PROMPT"] = prompt or "{text}"
        cfg["LLM"]["TASKS"]["editor"]["ENABLED"] = True
        cfg["LLM"]["TASKS"]["editor"]["PROVIDER"] = selected_provider
        cfg["LLM"]["TASKS"]["editor"]["COUNT_MIN"] = count_min
        cfg["LLM"]["TASKS"]["editor"]["COUNT_MAX"] = count_max
        save_config(cfg)

        self.log_signal.ai_log_message.emit("开始处理流程...")
        self.log_signal.ai_tasks_loaded.emit([])
        self.log_signal.ai_log_message.emit("【准备】复制原始文件...")
        import shutil
        copied_paths = []
        image_exts = (".png", ".jpg", ".jpeg", ".bmp", ".gif")
        for root, files in iter_files_limited(folder, max_depth=4):
            for file in files:
                name_check = file.lstrip()
                name_lower = file.lower()
                if (name_lower.endswith(".docx") or name_lower.endswith(image_exts)) and not name_check.startswith("~$") and not name_check.startswith("改 "):
                    original_path = os.path.join(root, file)
                    new_filename = f"改 {file}"
                    new_path = os.path.join(root, new_filename)
                    try:
                        shutil.copy2(original_path, new_path)
                        copied_paths.append(new_path)
                        self.log_signal.ai_log_message.emit(f"  {new_filename}")
                    except Exception as e:
                        self.log_signal.ai_log_message.emit(f"  {file} 复制失败: {e}")

        if not copied_paths:
            self.log_signal.ai_tasks_loaded.emit([])
            self.log_signal.ai_log_message.emit("未找到需要处理的文件")
            return

        self.log_signal.ai_tasks_loaded.emit(copied_paths)

        enabled_tasks = sorted([(t["id"], t["order"]) for t in self.task_config if t["enabled"]], key=lambda x: x[1])
        task_step_names = {"6": "DOC转DOCX", "1": "清除空格", "AI": "AI改作文", "2": "添加标签", "3": "格式化", "5": "改作者"}

        def update_step(task_path, step_name):
            self.log_signal.ai_task_status.emit(task_path, "running", step_name, "", f"开始：{step_name}")

        try:
            for task_id, _ in enabled_tasks:
                step_name = task_step_names.get(task_id, task_id)
                self.log_signal.ai_log_message.emit(f"【{step_name}】开始...")
                # 更新所有任务的当前步骤
                for task_path in copied_paths:
                    update_step(task_path, step_name)

                if task_id == "6":
                    self._convert_docs(folder)
                elif task_id == "1":
                    self._clear_spaces(folder)
                elif task_id == "AI":
                    self._process_ai(folder, api_key, base_url, prompt, count_min=count_min, count_max=count_max)
                elif task_id == "2":
                    self._add_labels(folder)
                elif task_id == "3":
                    self._format_docs(folder)
                elif task_id == "5":
                    self._set_author(folder)

            # 更新完成状态和修改后字数
            for task_path in copied_paths:
                output_path = self._ai_output_docx_path(task_path)
                after_count = self._count_docx_file_chars(output_path)
                self.log_signal.ai_task_status.emit(task_path, "done", "完成", after_count, "处理完成")

            self.log_signal.ai_log_message.emit("所有流程完成！")
        except Exception as e:
            for task_path in copied_paths:
                self.log_signal.ai_task_status.emit(task_path, "failed", "", "", f"失败: {e}")
            self.log_signal.ai_log_message.emit(f"处理失败：{e}")
            self.log_signal.ai_log_message.emit(_format_exception_for_log(e))
            traceback.print_exc()

    def _convert_docs(self, folder):
        import subprocess
        for root, files in iter_files_limited(folder, max_depth=4):
            for file in files:
                name_lower = file.lower()
                if name_lower.endswith(".doc") and not file.startswith("~$"):
                    doc_path = os.path.join(root, file)
                    try:
                        subprocess.run(["soffice", "--headless", "--convert-to", "docx", doc_path, "--outdir", root], capture_output=True, timeout=30)
                        base_name = os.path.splitext(os.path.basename(doc_path))[0]
                        new_path = os.path.join(root, base_name + ".docx")
                        if os.path.exists(new_path):
                            os.remove(doc_path)
                            self.log_signal.ai_log_message.emit(f"  {base_name}")
                    except Exception as e:
                        self.log_signal.ai_log_message.emit(f"  {file}: {e}")

    def _clear_spaces(self, folder):
        for root, files in iter_files_limited(folder, max_depth=4):
            for file in files:
                if file.lower().endswith(".docx") and not file.startswith("~$") and not file.startswith("改 "):
                    continue
                if file.lower().endswith(".docx") and not file.startswith("~$"):
                    try:
                        doc = Document(os.path.join(root, file))
                        for para in doc.paragraphs:
                            for run in para.runs:
                                run.text = run.text.strip()
                        doc.save(os.path.join(root, file))
                        self.log_signal.ai_log_message.emit(f"  {file}")
                    except Exception as e:
                        self.log_signal.ai_log_message.emit(f"  {file}: {e}")

    def _process_ai(self, folder, api_key, base_url, prompt_template, count_min=None, count_max=None):
        if not prompt_template:
            prompt_template = "下面是一篇中文文章，请你【只修改错别字和明显的识别错误】。\n要求：1. 不改变原意 2. 不润色文风 3. 不增删内容 4. 保持原有段落结构 5. 只输出修改后的完整文章正文\n"

        client = OpenAI(api_key=api_key, base_url=base_url)
        for root, files in iter_files_limited(folder, max_depth=4):
            for file in files:
                if not file.lower().endswith(".docx") or file.startswith("~$"):
                    continue
                if not file.startswith("改 "):
                    continue
                doc_path = os.path.join(root, file)
                try:
                    doc = Document(doc_path)
                    all_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip() and p.text.strip() not in ("修改前：", "修改后：")])
                    if not all_text.strip():
                        self.log_signal.ai_log_message.emit(f"  {file} (空文档)")
                        continue

                    original_count = count_chinese_characters(all_text)
                    if count_min is None or count_max is None:
                        default_min, default_max = determine_word_count_bounds(original_count)
                        count_min = count_min if count_min is not None else default_min
                        count_max = count_max if count_max is not None else default_max

                    for attempt in range(1, 5):
                        if "{text}" in prompt_template:
                            current_prompt = prompt_template.format(text=all_text)
                        else:
                            current_prompt = prompt_template + "\n\n" + all_text
                        current_prompt += f"\n\n请注意：这一次的修改后的正文总字数应控制在 {count_min} 到 {count_max} 之间"
                        if attempt > 1:
                            current_prompt += f"\n\n字数不符合规则，请重新修改并返回修改后的正文。只输出正文，不要解释。这次要求字数在 {count_min} 到 {count_max} 之间。"

                        self.log_signal.ai_log_message.emit(f"  {file} AI 第{attempt}次输出，正在检查字数...")
                        response = client.chat.completions.create(model="deepseek-chat", messages=[{"role": "system", "content": "你是一名严谨的中文校对助手"}, {"role": "user", "content": current_prompt}], temperature=0.1, stream=False)
                        result_text = response.choices[0].message.content.strip()
                        current_count = count_chinese_characters(result_text)
                        if count_min <= current_count <= count_max:
                            self.log_signal.ai_log_message.emit(f"  {file} 字数符合：{current_count}（目标 {count_min}-{count_max}）")
                            break
                        self.log_signal.ai_log_message.emit(f"  {file} 字数不合规：{current_count}，目标 {count_min}-{count_max}，正在重试...")
                    else:
                        raise RuntimeError(f"{file} AI 输出字数不符合要求")

                    last_para = doc.paragraphs[-1] if doc.paragraphs else None
                    if last_para:
                        if last_para.runs:
                            last_para.runs[-1].add_break(WD_BREAK.PAGE)
                        else:
                            last_para.add_run().add_break(WD_BREAK.PAGE)
                    para_modify = doc.add_paragraph("修改后：")
                    para_modify.paragraph_format.first_line_indent = Cm(0.74)
                    para_modify.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
                    para_modify.paragraph_format.line_spacing = Pt(12)
                    for line in result_text.split("\n"):
                        if line.strip():
                            p = doc.add_paragraph(line.strip())
                            p.paragraph_format.first_line_indent = Cm(0.74)
                            p.paragraph_format.space_before = Pt(0)
                            p.paragraph_format.space_after = Pt(0)
                            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
                            p.paragraph_format.line_spacing = Pt(12)
                    doc.save(doc_path)
                    self.log_signal.ai_log_message.emit(f"  {file}")
                except Exception as e:
                    self.log_signal.ai_log_message.emit(f"  {file}: {e}")
                    if _is_bad_marshal_error(e):
                        _clear_project_bytecode_caches()
                        self.log_signal.ai_log_message.emit(_format_exception_for_log(e))

    def _add_labels(self, folder):
        for root, files in iter_files_limited(folder, max_depth=4):
            for file in files:
                if not file.lower().endswith(".docx") or file.startswith("~$") or not file.startswith("改 "):
                    continue
                doc_path = os.path.join(root, file)
                try:
                    doc = Document(doc_path)
                    if doc.paragraphs:
                        last_para = doc.paragraphs[-1]
                        has_modify = last_para.text.strip() == "修改后：" or (len(doc.paragraphs) > 1 and doc.paragraphs[-2].text.strip() == "修改后：")
                        if doc.paragraphs[0].text.strip() != "修改前：":
                            doc.paragraphs[0].insert_paragraph_before("修改前：")
                        if not has_modify:
                            last_para = doc.paragraphs[-1]
                            if last_para.runs:
                                last_para.runs[-1].add_break(WD_BREAK.PAGE)
                            else:
                                last_para.add_run().add_break(WD_BREAK.PAGE)
                            para = doc.add_paragraph("修改后：")
                            para.paragraph_format.first_line_indent = Cm(0.74)
                            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
                            para.paragraph_format.line_spacing = Pt(12)
                    doc.save(doc_path)
                    self.log_signal.ai_log_message.emit(f"  {file}")
                except Exception as e:
                    self.log_signal.ai_log_message.emit(f"  {file}: {e}")

    def _format_docs(self, folder):
        for root, files in iter_files_limited(folder, max_depth=4):
            for file in files:
                if not file.lower().endswith(".docx") or file.startswith("~$") or not file.startswith("改 "):
                    continue
                doc_path = os.path.join(root, file)
                try:
                    doc = Document(doc_path)
                    style = doc.styles["Normal"]
                    style.font.name = "宋体"
                    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    style.font.size = Pt(12)
                    for para in doc.paragraphs:
                        para.paragraph_format.first_line_indent = Cm(0.74)
                        para.paragraph_format.space_before = Pt(0)
                        para.paragraph_format.space_after = Pt(0)
                        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
                        para.paragraph_format.line_spacing = Pt(12)
                    doc.save(doc_path)
                    self.log_signal.ai_log_message.emit(f"  {file}")
                except Exception as e:
                    self.log_signal.ai_log_message.emit(f"  {file}: {e}")

    def _set_author(self, folder):
        for root, files in iter_files_limited(folder, max_depth=4):
            for file in files:
                if not file.lower().endswith(".docx") or file.startswith("~$") or not file.startswith("改 "):
                    continue
                doc_path = os.path.join(root, file)
                try:
                    doc = Document(doc_path)
                    doc.core_properties.author = "思睿教育_美丽可爱的尹老师"
                    doc.save(doc_path)
                    self.log_signal.ai_log_message.emit(f"  {file}")
                except Exception as e:
                    self.log_signal.ai_log_message.emit(f"  {file}: {e}")


# ===================== Main =====================
if __name__ == "__main__":
    try:
        _write_startup_log("main: start")
        app = QApplication(sys.argv)
        _write_startup_log("main: QApplication created")
        app.setStyle("Fusion")
        window = MainWindow()
        _write_startup_log("main: MainWindow created")
        window.show()
        _write_startup_log(f"main: window shown visible={window.isVisible()}")
        exit_code = app.exec()
        _write_startup_log(f"main: app.exec returned {exit_code}")
        sys.exit(exit_code)
    except Exception:
        _write_crash_log(*sys.exc_info())
        raise


